#!/usr/bin/env python3
"""
=============================================================================
  Solar & Battery Economic Modeller — Synergy WA (Perth / SWIS)
  Effective tariff rates: 1 July 2025
  Shading model: seasonal + diurnal (user-specified ~40% annual avg)
=============================================================================
"""

import argparse, sys, warnings
from pathlib import Path
from itertools import product as itertools_product

try:
    import numpy as np
    import pandas as pd
    import matplotlib; matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.dates as mdates
except ImportError as exc:
    sys.exit(f"Missing: {exc}\n  pip install numpy pandas matplotlib openpyxl")

warnings.filterwarnings("ignore")

# =============================================================================
#  USER CONFIGURATION — edit this section to customise your analysis
# =============================================================================

# ─── DATA FILE ───────────────────────────────────────────────────────────────
DEFAULT_DATA_PATH = "/Users/jonathanking/Python/solar_analysis_output/MyHalfHourlyMeterData-2024.xlsx"

# ── How to read these quotes:
# Enter the GROSS installed price (before rebates) — code deducts rebates automatically.
# STC solar rebate is usually already in the installer's quoted price, so don't add it back.

# ─── QUOTES (label, solar_kW, battery_kWh, inverter_kW, total_cost_AUD) ─────
# IMPORTANT: total_cost_AUD must be the GROSS installed price BEFORE deducting
# the federal battery rebate and state battery rebate. The code deducts both
# automatically. STC solar rebates are typically already applied upfront by the
# installer, so the price you enter should already reflect that discount.
QUOTES = [
    ("Base Case — Grid Only",            0.0,   0.0,  0.0,      0),  # current situation, no solar/battery
    ("12.32 kW Solar + 16 kWh Battery", 12.3,  16.0, 10.0, 25_180),
    ("6.6 kW Solar + 26.3 kWh Battery",  6.3,  26.3, 10.0, 25_280),
    ("6.6 kW Solar + 17.5 kWh Battery",  6.3,  17.5, 10.0, 21_370),
    #("12 kW Solar + 14 kWh Battery",  12.0,  14.0,  10.0, 18_490),
    
]


# ─── ANALYSIS HORIZON ────────────────────────────────────────────────────────
ANALYSIS_YEARS = 20

# ─── TARIFF CONSTANTS (1 Jul 2025) ───────────────────────────────────────────
A1_SUPPLY   = 116.0505 / 100   # $/day — fixed daily connection fee on the A1 Flat tariff
A1_RATE     =  32.3719 / 100   # $/kWh — flat import rate for every kWh on A1 Flat (same all day)
MS_SUPPLY   = 116.0505 / 100   # $/day — fixed daily connection fee on the Midday Saver tariff
MS_SOP      =   8.6151 / 100   # $/kWh — super off-peak (9am–3pm): very cheap, the key charging window
MS_OP       =  22.1777 / 100   # $/kWh — off-peak (9pm–9am): standard overnight rate
MS_PEAK     =  53.8446 / 100   # $/kWh — peak (3pm–9pm): expensive; avoid importing here
DEBS_PEAK   =  10.0    / 100   # $/kWh — DEBS export credit rate during the 3–9pm peak window
DEBS_OP     =   2.0    / 100   # $/kWh — DEBS export credit rate at all other times (very low)

TARIFF_ESC  = 0.025   # electricity prices assumed to rise 2.5% per year over the analysis horizon
DEBS_DECL   = 0.050   # DEBS export credit rate falls 5% per year (government policy direction)

# ─── DISCOUNT RATE ───────────────────────────────────────────────────────────
# Used to compute NPV-adjusted ("discounted") payback: the year when cumulative
# PV-discounted savings first exceed the upfront net cost.
OPPORTUNITY_RATE = 0.08   # 8% p.a. — discount / hurdle rate for NPV calculations

# ─── BATTERY / SOLAR PARAMETERS ──────────────────────────────────────────────
BAT_RTE    = 0.90    # round-trip efficiency: 90% of energy put in can be taken back out (√ applied each way)
BAT_DOD    = 0.90    # depth of discharge: only 90% of nameplate capacity is usable (protects battery life)
BAT_DEG    = 0.02    # battery capacity degrades 2% per year (so a 13.5 kWh battery holds ~11 kWh by year 10)
SOL_DEG    = 0.005   # solar panels lose 0.5% of their output per year due to cell ageing
SYS_EFF    = 0.80    # inverter + wiring losses: only 80% of DC solar generation reaches the meter

# ─── SHADING FACTORS (fraction of unshaded output — Southern Hemisphere) ─────
# Adjust these to match your site. 1.0 = fully unshaded; lower = more shading.
# These defaults represent a typical Perth suburban roof with moderate tree/fence shading.
# Fully unshaded north-facing roof: use 0.90–0.95 for summer, 0.80–0.85 for winter.
# Heavily shaded site: use 0.50–0.60 for summer, 0.40–0.50 for winter.
SHADE_SUMMER = 0.60  # Dec–Feb: sun high, shading mainly confined to early morning/late afternoon
SHADE_AUTUMN = 0.50  # Mar–May: transitional; moderate shadow angles through the day
SHADE_WINTER = 0.40  # Jun–Aug: sun low, some shading across more of the day
SHADE_SPRING = 0.50  # Sep–Nov: recovering toward summer; similar to autumn

# ─── SOLAR CALIBRATION — Perth-specific ──────────────────────────────────────
# Perth latitude and unshaded specific yield.  SOLAR_K is computed at startup from
# these so that a 1 kWp system with no shading (all SHADE_* = 1.0) produces exactly
# PERTH_SPECIFIC_YIELD kWh/year.  The physically-based sin(elevation) shape is used
# so the profile is accurate — not a generic bell curve.
PERTH_LAT_DEG       = -31.95   # Perth latitude (degrees, negative = south)
PERTH_SPECIFIC_YIELD = 1650    # kWh/kWp/year — unshaded, post SYS_EFF, north-facing Perth

# ─── REBATES (1 May 2026 rates) ──────────────────────────────────────────────
STATE_REBATE_FLAT = 1_300.0   # Synergy WA Residential Battery Scheme ($130/kWh, cap $1,300)


def fed_battery_rebate(bat_kwh: float) -> float:
    """Cheaper Home Batteries Program — tiered STC rebate from 1 May 2026.
    Tier 1:  0–14 kWh  @ ~$252/kWh
    Tier 2: 14–28 kWh  @ ~$151/kWh
    Tier 3: 28–50 kWh  @  ~$38/kWh
    Over 50 kWh: no rebate.
    """
    tier1 = min(bat_kwh, 14.0) * 252.0
    tier2 = max(0.0, min(bat_kwh, 28.0) - 14.0) * 151.0
    tier3 = max(0.0, min(bat_kwh, 50.0) - 28.0) * 38.0
    return tier1 + tier2 + tier3

# =============================================================================

# ─── VALIDATION ──────────────────────────────────────────────────────────────
class ValidationError(Exception):
    """Raised when input data or configuration fails validation."""
    pass


def _warn(msg: str):
    """Print a yellow warning message."""
    print(f"  ⚠ WARNING: {msg}")


def _info(msg: str):
    print(f"  ℹ {msg}")


def validate_file_path(path) -> Path:
    """Check that the data file exists and looks like a supported format."""
    p = Path(path)
    if not p.exists():
        raise ValidationError(f"Data file not found: {path}")
    if not p.is_file():
        raise ValidationError(f"Path is not a file: {path}")
    suffix = p.suffix.lower()
    if suffix not in (".xlsx", ".xls", ".csv", ".tsv"):
        raise ValidationError(
            f"Unsupported file type '{suffix}'. Expected .xlsx, .xls, .csv, or .tsv."
        )
    if p.stat().st_size == 0:
        raise ValidationError(f"Data file is empty: {path}")
    if p.stat().st_size > 200 * 1024 * 1024:    # 200 MB
        _warn(f"File is unusually large ({p.stat().st_size/1e6:.1f} MB). "
              f"Half-hourly data for one year is typically <2 MB.")
    return p


def validate_consumption_data(df: pd.DataFrame) -> pd.DataFrame:
    """
    Validate the loaded half-hourly consumption DataFrame.

    Checks: column presence, row count, datetime sanity, gaps, duplicates,
    plausible kWh values. Returns the (possibly-cleaned) DataFrame.
    """
    if df is None or len(df) == 0:
        raise ValidationError("Consumption DataFrame is empty.")

    # ── Required columns
    required = {"datetime", "consumption_kwh"}
    missing  = required - set(df.columns)
    if missing:
        raise ValidationError(f"Missing required columns: {missing}")

    # ── datetime column type
    if not pd.api.types.is_datetime64_any_dtype(df["datetime"]):
        raise ValidationError("'datetime' column is not a datetime type.")

    # ── consumption column type
    if not pd.api.types.is_numeric_dtype(df["consumption_kwh"]):
        raise ValidationError("'consumption_kwh' column is not numeric.")

    # ── Check for NaN/null
    n_null_dt   = df["datetime"].isna().sum()
    n_null_kwh  = df["consumption_kwh"].isna().sum()
    if n_null_dt:
        raise ValidationError(f"{n_null_dt} rows have null datetime values.")
    if n_null_kwh:
        _warn(f"{n_null_kwh} rows have null kWh values — these will be treated as 0.")
        df["consumption_kwh"] = df["consumption_kwh"].fillna(0)

    # ── Negative consumption (impossible for a passive load meter)
    n_neg = (df["consumption_kwh"] < 0).sum()
    if n_neg:
        _warn(f"{n_neg} rows have negative kWh values — clamping to 0. "
              f"This may indicate net-metered exports already in your data.")
        df.loc[df["consumption_kwh"] < 0, "consumption_kwh"] = 0

    # ── Suspiciously high half-hourly values
    # A typical residential connection is 63A single-phase = ~14.5 kW
    # 14.5 kW × 0.5 h = 7.25 kWh max per half-hour interval
    # Using 10 kWh/30min (≈ 20 kW) as a conservative outlier threshold
    n_high = (df["consumption_kwh"] > 10).sum()
    if n_high:
        max_v = df["consumption_kwh"].max()
        _warn(f"{n_high} half-hour intervals exceed 10 kWh "
              f"(max: {max_v:.2f} kWh). Plausible for a 3-phase property "
              f"or EV charging — verify this is correct.")

    # ── Duplicate timestamps
    n_dup = df["datetime"].duplicated().sum()
    if n_dup:
        _warn(f"{n_dup} duplicate timestamps found — keeping first occurrence.")
        df = df.drop_duplicates(subset="datetime", keep="first").reset_index(drop=True)

    # ── Sort
    df = df.sort_values("datetime").reset_index(drop=True)

    # ── Date range coverage
    span_days = (df["datetime"].iloc[-1] - df["datetime"].iloc[0]).days + 1
    if span_days < 30:
        raise ValidationError(
            f"Only {span_days} days of data — need at least 30 days for any "
            f"meaningful analysis, ideally a full year (365 days)."
        )
    if span_days < 350:
        _warn(f"Only {span_days} days of data — annual figures will be "
              f"extrapolated. Best results come from a full year.")
    elif span_days > 380:
        _warn(f"Data spans {span_days} days — more than one year. "
              f"Analysis will use all data; figures may not represent a single year.")

    # ── Half-hourly cadence check
    deltas = df["datetime"].diff().dropna()
    expected = pd.Timedelta(minutes=30)
    n_off    = (deltas != expected).sum()
    if n_off > 0:
        # Identify gaps and irregular intervals
        gaps = deltas[deltas > expected]
        n_gaps = len(gaps)
        if n_gaps:
            total_missing = int(((gaps - expected).dt.total_seconds() / 1800).sum())
            _warn(f"{n_gaps} gaps in the time series totalling ~{total_missing} "
                  f"missing half-hour intervals. Annual totals may be slightly understated.")

        n_short = (deltas < expected).sum()
        if n_short:
            _warn(f"{n_short} intervals shorter than 30 min — possible duplicates or "
                  f"daylight-saving artefacts.")

    # ── Check for full coverage when length is right
    expected_rows_year     = 365 * 48
    expected_rows_leap     = 366 * 48
    if len(df) not in (expected_rows_year, expected_rows_leap):
        if abs(len(df) - expected_rows_year) <= 48 * 7:   # within a week
            _info(f"{len(df):,} rows loaded (expected {expected_rows_year:,} for "
                  f"a non-leap year or {expected_rows_leap:,} for a leap year).")
        else:
            _warn(f"{len(df):,} rows loaded — expected ~{expected_rows_year:,} "
                  f"for a full year of half-hourly data.")

    # ── Sanity check on totals
    total_kwh = df["consumption_kwh"].sum()
    avg_daily = total_kwh / max(span_days, 1)
    if avg_daily < 1.0:
        _warn(f"Average daily consumption is only {avg_daily:.2f} kWh — "
              f"check the kWh column values aren't in Wh by mistake.")
    elif avg_daily > 200:
        _warn(f"Average daily consumption is {avg_daily:.0f} kWh — "
              f"very high for a residential property. Verify this is the right meter.")

    if total_kwh == 0:
        raise ValidationError("Total consumption is 0 kWh — data appears empty.")

    return df


def validate_quotes(quotes: list) -> list:
    """Validate the list of system quotes; raise on hard errors, warn on soft."""
    if not quotes:
        raise ValidationError("QUOTES list is empty — add at least one quote.")

    cleaned = []
    seen_labels = set()

    for i, q in enumerate(quotes):
        # Tuple shape
        if not isinstance(q, (tuple, list)) or len(q) != 5:
            raise ValidationError(
                f"Quote {i+1}: expected a 5-tuple "
                f"(label, solar_kW, battery_kWh, inverter_kW, cost_AUD). Got: {q!r}"
            )

        label, solar_kw, bat_kwh, inv_kw, cost = q

        # Label
        if not isinstance(label, str) or not label.strip():
            raise ValidationError(f"Quote {i+1}: label must be a non-empty string.")
        if label in seen_labels:
            _warn(f"Quote {i+1}: duplicate label '{label}'. "
                  f"Charts and tables may be ambiguous.")
        seen_labels.add(label)

        is_base_case = (solar_kw == 0 and bat_kwh == 0 and inv_kw == 0 and cost == 0)

        # Numeric fields
        for name, val, lo, hi in [
            ("solar_kW",     solar_kw, 0.5, 50),
            ("battery_kWh",  bat_kwh,  0,   100),
            ("inverter_kW",  inv_kw,   1,   30),
            ("cost_AUD",     cost,     1_000, 200_000),
        ]:
            if not isinstance(val, (int, float)) or val != val:    # NaN check
                raise ValidationError(
                    f"Quote '{label}': {name} must be a number. Got: {val!r}"
                )
            if val < 0:
                raise ValidationError(f"Quote '{label}': {name} cannot be negative.")
            if not is_base_case:
                if val < lo:
                    _warn(f"Quote '{label}': {name}={val} is unusually low (expected ≥ {lo}).")
                if val > hi:
                    _warn(f"Quote '{label}': {name}={val} is unusually high (expected ≤ {hi}).")

        # Cross-field sanity
        # Inverter typically 0.7–1.3× solar size for solar-only AC sizing.
        # With a battery, the inverter handles both PV + battery so it can be
        # smaller than solar (PV is "clipped") but should be ≥ 0.5× solar.
        if solar_kw > 0:
            ratio = inv_kw / solar_kw
            if ratio < 0.5:
                _warn(f"Quote '{label}': inverter ({inv_kw}kW) is < 50% of solar "
                      f"({solar_kw}kW) — significant clipping losses likely.")
            elif ratio > 1.5:
                _warn(f"Quote '{label}': inverter ({inv_kw}kW) is > 150% of solar "
                      f"({solar_kw}kW) — oversized for the array.")

        # Cost-per-kW sanity (typical Perth 2025: $1.0–2.5k per installed kW solar,
        # $0.7–1.3k per kWh battery before rebates)
        if bat_kwh > 0:
            implied = solar_kw * 1500 + bat_kwh * 1000
        else:
            implied = solar_kw * 1500
        if cost > 0 and implied > 0:
            ratio = cost / implied
            if ratio < 0.5:
                _warn(f"Quote '{label}': ${cost:,.0f} looks low for "
                      f"{solar_kw}kW + {bat_kwh}kWh (~${implied:,.0f} expected). "
                      f"Verify cost includes installation.")
            elif ratio > 2.0:
                _warn(f"Quote '{label}': ${cost:,.0f} looks high for "
                      f"{solar_kw}kW + {bat_kwh}kWh (~${implied:,.0f} expected).")

        cleaned.append((label, float(solar_kw), float(bat_kwh),
                        float(inv_kw), float(cost)))
    return cleaned


def validate_constants():
    """Sanity-check the tariff and system parameter constants."""
    errors = []

    # Tariff rates must be positive and within sensible bounds
    for name, val, lo, hi in [
        ("A1_RATE",   A1_RATE,   0.10, 1.00),
        ("A1_SUPPLY", A1_SUPPLY, 0.50, 5.00),
        ("MS_SOP",    MS_SOP,    0.01, 0.30),
        ("MS_OP",     MS_OP,     0.10, 0.50),
        ("MS_PEAK",   MS_PEAK,   0.30, 1.00),
        ("DEBS_PEAK", DEBS_PEAK, 0.01, 0.50),
        ("DEBS_OP",   DEBS_OP,   0.00, 0.30),
    ]:
        if not (lo <= val <= hi):
            errors.append(f"{name}={val:.4f} outside expected range [{lo}, {hi}]")

    # Midday Saver structure: super-off-peak < off-peak < peak
    if not (MS_SOP < MS_OP < MS_PEAK):
        errors.append(f"Midday Saver rates not monotonic: "
                      f"SOP={MS_SOP} OP={MS_OP} PEAK={MS_PEAK}")

    # DEBS peak should pay more than DEBS off-peak
    if DEBS_PEAK <= DEBS_OP:
        errors.append(f"DEBS peak rate ({DEBS_PEAK}) should exceed off-peak ({DEBS_OP})")

    # Battery / system parameters must be 0 < x ≤ 1 for fractions
    for name, val in [("BAT_RTE", BAT_RTE), ("BAT_DOD", BAT_DOD), ("SYS_EFF", SYS_EFF)]:
        if not (0 < val <= 1):
            errors.append(f"{name}={val} must be in (0, 1].")

    # Degradation must be 0–10% p.a.
    for name, val in [("BAT_DEG", BAT_DEG), ("SOL_DEG", SOL_DEG)]:
        if not (0 <= val <= 0.10):
            errors.append(f"{name}={val} should be 0–10% p.a.")

    # Escalation rates 0–15% p.a.
    for name, val in [("TARIFF_ESC", TARIFF_ESC), ("DEBS_DECL", DEBS_DECL)]:
        if not (-0.15 <= val <= 0.15):
            errors.append(f"{name}={val} outside ±15% p.a.")

    if not (1 <= ANALYSIS_YEARS <= 50):
        errors.append(f"ANALYSIS_YEARS={ANALYSIS_YEARS} should be 1–50.")

    # Rebates non-negative
    if STATE_REBATE_FLAT < 0:
        errors.append("Rebate values must be non-negative.")

    # Shading factors must be between 0 and 1
    for name, val in [("SHADE_SUMMER", SHADE_SUMMER), ("SHADE_AUTUMN", SHADE_AUTUMN),
                      ("SHADE_WINTER", SHADE_WINTER), ("SHADE_SPRING", SHADE_SPRING)]:
        if not (0.0 < val <= 1.0):
            errors.append(f"{name}={val} must be between 0 and 1.")
    if not (SHADE_WINTER <= SHADE_SPRING <= SHADE_AUTUMN <= SHADE_SUMMER):
        errors.append("Expected SHADE_WINTER ≤ SHADE_SPRING ≤ SHADE_AUTUMN ≤ SHADE_SUMMER "
                      "(Southern Hemisphere seasonal ordering).")

    if errors:
        raise ValidationError(
            "Configuration constants failed validation:\n   - "
            + "\n   - ".join(errors)
        )


# ─── SHADING MODEL ───────────────────────────────────────────────────────────
# ─── SOLAR GEOMETRY & CALIBRATION ────────────────────────────────────────────
# Compute SOLAR_K once at startup: the scalar that makes a 1 kWp unshaded system
# in Perth produce exactly PERTH_SPECIFIC_YIELD kWh/year when driven by sin(elevation).
_lat = np.radians(PERTH_LAT_DEG)

def _solar_k():
    lat = _lat
    total = 0.0
    for doy in range(1, 366):
        decl = np.radians(-23.45 * np.cos(2*np.pi*(doy+10)/365))
        ha_arr = np.radians(15.0 * ((np.arange(48)+0.5)/2.0 - 12.0))
        sin_e = np.sin(lat)*np.sin(decl) + np.cos(lat)*np.cos(decl)*np.cos(ha_arr)
        total += np.maximum(0.0, sin_e).sum() * 0.5  # 0.5 h per slot
    return PERTH_SPECIFIC_YIELD / (total * SYS_EFF)

SOLAR_K = _solar_k()


def shading_factor(slot, doy):
    """Shading fraction for display heatmap: seasonal × diurnal taper (lower sun = more shade)."""
    m = (pd.Timestamp("2024-01-01") + pd.Timedelta(days=int(doy)-1)).month
    if m in (12,1,2):   seasonal = SHADE_SUMMER
    elif m in (3,4,5):  seasonal = SHADE_AUTUMN
    elif m in (6,7,8):  seasonal = SHADE_WINTER
    else:               seasonal = SHADE_SPRING
    # Diurnal taper: shade is worse near sunrise/sunset (low sun angle = longer shadows).
    # sin(elevation) already goes to zero at horizon, so this taper is visual-only.
    decl = np.radians(-23.45 * np.cos(2*np.pi*(doy+10)/365))
    ha   = np.radians(15.0 * (slot/2.0 - 12.0))
    sin_e = np.sin(_lat)*np.sin(decl) + np.cos(_lat)*np.cos(decl)*np.cos(ha)
    if sin_e <= 0:
        return 0.0
    # Taper from 100% at peak sun to 70% near horizon — shading is angle-dependent
    return seasonal * (0.70 + 0.30 * sin_e)


def solar_profile(solar_kw, doy):
    """
    48-element array (kWh per half-hour) using Perth solar geometry.
    Shape = sin(solar elevation angle) — physically correct for Perth latitude.
    Magnitude calibrated so 1 kWp × SHADE=1.0 → PERTH_SPECIFIC_YIELD kWh/year.
    """
    decl  = np.radians(-23.45 * np.cos(2*np.pi*(doy+10)/365))
    slots = np.arange(48)
    hours = (slots + 0.5) / 2.0          # mid-point of each 30-min slot
    ha    = np.radians(15.0*(hours-12.0))
    sin_e = np.sin(_lat)*np.sin(decl) + np.cos(_lat)*np.cos(decl)*np.cos(ha)
    irrad = np.maximum(0.0, sin_e)        # zero when sun is below horizon

    m = (pd.Timestamp("2024-01-01") + pd.Timedelta(days=int(doy)-1)).month
    if m in (12,1,2):   sf = SHADE_SUMMER
    elif m in (3,4,5):  sf = SHADE_AUTUMN
    elif m in (6,7,8):  sf = SHADE_WINTER
    else:               sf = SHADE_SPRING

    # kWh = kW × irrad_fraction × calibration × system_efficiency × 0.5h × site_shading
    return solar_kw * irrad * SOLAR_K * SYS_EFF * 0.5 * sf


# ─── RATE HELPERS ─────────────────────────────────────────────────────────────
def ms_rate(slot):
    h = slot/2.0
    if 9 <= h < 15:  return MS_SOP
    if 15 <= h < 21: return MS_PEAK
    return MS_OP

def debs_rate(slot): return DEBS_PEAK if 30<=slot<42 else DEBS_OP


# ─── SIMULATION ───────────────────────────────────────────────────────────────
def add_solar(raw_df, solar_kw):
    """Return a copy of raw_df with solar_kwh column attached for a given system size."""
    sol_cache = {}
    def get_solar(row):
        doy = int(row["doy"])
        if doy not in sol_cache:
            sol_cache[doy] = solar_profile(solar_kw, doy)
        return sol_cache[doy][int(row["slot"])]
    df = raw_df.copy()
    df["solar_kwh"] = df.apply(get_solar, axis=1)
    return df


def simulate(df, solar_kw, bat_kwh, inv_kw, tariff, yr_offset=0):
    esc      = (1+TARIFF_ESC)**yr_offset
    debs_esc = max(0.005, 1-DEBS_DECL*yr_offset)
    usable   = bat_kwh * BAT_DOD
    soc      = usable * 0.5
    max_chg  = min(inv_kw, solar_kw)

    n = len(df)
    g_imp = np.zeros(n); g_exp = np.zeros(n)
    b_chg = np.zeros(n); b_dis = np.zeros(n)
    s_slf = np.zeros(n); soc_a = np.zeros(n)

    solar_arr = df["solar_kwh"].values
    load_arr  = df["consumption_kwh"].values
    slot_arr  = df["slot"].values

    # ── Battery dispatch logic (runs for every 30-minute interval in order):
    # If solar > load:  charge battery from surplus, export remainder
    # If solar < load:  discharge battery to cover shortfall, import remainder
    # SoC = state of charge (kWh stored), clamped to [0, usable capacity]
    # All energy through the battery loses sqrt(RTE) on the way in and out
    for i in range(n):
        load  = load_arr[i]
        solar = solar_arr[i]
        net   = solar - load
        if net >= 0:
            s_slf[i] = load
            chg = min(net, max_chg*0.5, (usable-soc)/BAT_RTE**0.5)  # how much surplus solar can the battery absorb this interval?
            soc = min(soc + chg*BAT_RTE**0.5, usable)
            g_exp[i] = max(0, net - chg)
            b_chg[i] = chg
        else:
            s_slf[i] = solar
            dis = min(-net, inv_kw*0.5, soc)  # how much stored energy can cover the shortfall?
            soc = max(0, soc - dis)
            g_imp[i] = max(0, -net - dis*BAT_RTE**0.5)
            b_dis[i] = dis
        soc_a[i] = soc

    df = df.copy()
    df["g_imp"] = g_imp; df["g_exp"] = g_exp
    df["b_chg"] = b_chg; df["b_dis"] = b_dis
    df["s_slf"] = s_slf; df["soc"]   = soc_a

    n_days = int(df["date"].nunique()) if "date" in df.columns else 365
    if tariff == "A1 Flat":
        ic = (g_imp * A1_RATE * esc).sum()
        sc = A1_SUPPLY * n_days * esc
    else:
        rates = np.vectorize(ms_rate)(slot_arr.astype(int)) * esc
        ic    = (g_imp * rates).sum()
        sc    = MS_SUPPLY * n_days * esc

    ec = (g_exp * np.vectorize(debs_rate)(slot_arr.astype(int)) * debs_esc).sum()

    tot_load  = load_arr.sum()
    tot_solar = solar_arr.sum()
    tot_self  = s_slf.sum()
    tot_bdis  = b_dis.sum() * BAT_RTE**0.5
    self_suf  = (tot_self+tot_bdis)/tot_load*100 if tot_load else 0

    return dict(df=df, tariff=tariff, solar_kw=solar_kw, bat_kwh=bat_kwh,
                label="", annual_import_kwh=g_imp.sum(), annual_export_kwh=g_exp.sum(),
                annual_solar_kwh=tot_solar, annual_load_kwh=tot_load,
                import_cost=ic, supply_cost=sc, export_credit=ec,
                net_cost=ic+sc-ec, self_suf_pct=self_suf)


def simulate_grid_arbitrage(df, bat_kwh, inv_kw, yr_offset=0):
    """Battery-only grid arbitrage on Midday Saver — no solar.

    Strategy: charge the battery from grid during the cheap super off-peak window
    (9 am–3 pm, 8.6¢/kWh) and discharge it during the expensive peak window
    (3–9 pm, 53.8¢/kWh).  Off-peak hours the battery is idle and grid covers load.
    The net benefit is ~45¢/kWh saved (before round-trip efficiency losses).
    """
    # ── Arbitrage strategy:
    # SOP (9am–3pm):  pull extra power from grid to fill the battery cheaply
    # Peak (3–9pm):   discharge battery to cover load, avoiding expensive peak import
    # Off-peak:       battery idle; grid covers load at off-peak rate
    # Net benefit per kWh shifted ≈ peak_rate - SOP_rate/RTE ≈ 44¢/kWh
    esc  = (1 + TARIFF_ESC) ** yr_offset
    rte  = BAT_RTE ** 0.5          # charging AND discharging half-efficiency
    b_cap = bat_kwh * BAT_DOD      # bat_kwh passed already degraded from payback

    n = len(df)
    load_arr = df["consumption_kwh"].values
    slot_arr = df["slot"].values

    g_imp = np.zeros(n); b_chg = np.zeros(n)
    b_dis = np.zeros(n); soc_a = np.zeros(n)

    soc = 0.0   # start empty each year

    for i in range(n):
        h    = slot_arr[i] / 2.0
        load = load_arr[i]

        if 9 <= h < 15:          # SOP — charge from grid
            space   = b_cap - soc
            chg_in  = min(inv_kw * 0.5, space / rte)   # grid draw to fill space
            chg_in  = max(0.0, chg_in)
            soc    += chg_in * rte
            b_chg[i] = chg_in
            g_imp[i] = load + chg_in                    # load + charging draw

        elif 15 <= h < 21:       # Peak — discharge to cover load
            # How much SoC to remove so that SoC*rte covers as much of load as possible?
            dis  = min(soc, inv_kw * 0.5)               # raw SoC reduction
            if dis * rte > load:                         # don't over-discharge
                dis = load / rte
            dis     = max(0.0, dis)
            soc    -= dis
            b_dis[i] = dis
            g_imp[i] = max(0.0, load - dis * rte)       # residual load from grid

        else:                    # Off-peak — battery idle
            g_imp[i] = load

        soc_a[i] = soc

    df = df.copy()
    df["solar_kwh"] = 0.0
    df["g_imp"] = g_imp;  df["g_exp"] = np.zeros(n)
    df["b_chg"] = b_chg;  df["b_dis"] = b_dis
    df["s_slf"] = np.zeros(n); df["soc"] = soc_a

    n_days = int(df["date"].nunique()) if "date" in df.columns else 365
    rates  = np.vectorize(ms_rate)(slot_arr.astype(int)) * esc
    ic     = (g_imp * rates).sum()
    sc     = MS_SUPPLY * n_days * esc

    tot_load = load_arr.sum()
    tot_bdis = (b_dis * rte).sum()
    self_suf = tot_bdis / tot_load * 100 if tot_load else 0

    return dict(df=df, tariff="Midday Saver", solar_kw=0.0, bat_kwh=bat_kwh,
                label="", annual_import_kwh=g_imp.sum(), annual_export_kwh=0.0,
                annual_solar_kwh=0.0, annual_load_kwh=tot_load,
                import_cost=ic, supply_cost=sc, export_credit=0.0,
                net_cost=ic + sc, self_suf_pct=self_suf)


def baseline(df, tariff):
    slots  = df["slot"].values
    load   = df["consumption_kwh"].values
    n_days = int(df["date"].nunique()) if "date" in df.columns else 365
    if tariff == "A1 Flat":
        return (load*A1_RATE).sum() + A1_SUPPLY*n_days
    return (load * np.vectorize(ms_rate)(slots.astype(int))).sum() + MS_SUPPLY*n_days


def payback(df, solar_kw, bat_kwh, inv_kw, cost, tariff, label):
    # ── 25-year payback model:
    # Year 0: pay net system cost (gross price minus government rebates)
    # Years 1–25: solar+battery output degrades slightly each year
    #             electricity prices rise at TARIFF_ESC % per year
    #             annual saving = (baseline cost that year) - (modelled cost that year)
    # Payback year = first year where cumulative savings exceed upfront cost
    fed  = min(fed_battery_rebate(bat_kwh), cost*0.4)
    state = STATE_REBATE_FLAT if bat_kwh>0 else 0
    net  = max(0, cost - fed - state)
    base = baseline(df, tariff)

    savings=[]; net_costs=[]; cum=[-net]; cum_disc=[-net]
    yr1_r = None
    for yr in range(1, ANALYSIS_YEARS+1):
        s_yr = solar_kw*(1-SOL_DEG)**yr
        b_yr = max(0, bat_kwh*(1-BAT_DEG)**yr)
        r    = simulate(df, s_yr, b_yr, inv_kw, tariff, yr_offset=yr)
        if yr == 1:
            yr1_r = r
        base_yr = base*(1+TARIFF_ESC)**yr
        sav  = base_yr - r["net_cost"]
        savings.append(sav); net_costs.append(r["net_cost"])
        cum.append(cum[-1]+sav)
        cum_disc.append(cum_disc[-1] + sav / (1+OPPORTUNITY_RATE)**yr)

    pb_yr      = 0 if net == 0 else next((yr for yr,cf in enumerate(cum[1:],1) if cf>=0), None)
    pb_yr_disc = 0 if net == 0 else next((yr for yr,cf in enumerate(cum_disc[1:],1) if cf>=0), None)
    npv        = cum_disc[-1]
    return dict(label=label, solar_kw=solar_kw, bat_kwh=bat_kwh, inv_kw=inv_kw,
                cost=cost, fed=fed, state=state, net=net, tariff=tariff,
                pb_yr=pb_yr, pb_yr_disc=pb_yr_disc, npv=npv,
                savings=savings, net_costs=net_costs, cum=cum, cum_disc=cum_disc,
                total_save=sum(savings),
                roi=sum(savings)/net*100 if net>0 else 0,
                yr1_save=savings[0] if savings else 0,
                self_suf_pct=yr1_r["self_suf_pct"] if yr1_r else 0.0,
                annual_export_kwh=yr1_r["annual_export_kwh"] if yr1_r else 0.0,
                annual_solar_kwh=yr1_r["annual_solar_kwh"] if yr1_r else 0.0)


def payback_arbitrage(raw_df, bat_kwh, inv_kw, cost, label):
    """25-year payback for a battery-only grid arbitrage scenario (Midday Saver)."""
    # Same structure as payback() but uses simulate_grid_arbitrage() each year.
    # No solar degradation — only battery capacity degrades over time.
    fed   = min(fed_battery_rebate(bat_kwh), cost * 0.4)
    state = STATE_REBATE_FLAT   # requires VPP connection
    net   = max(0.0, cost - fed - state)
    base  = baseline(raw_df, "Midday Saver")

    savings = []; net_costs = []; cum = [-net]; cum_disc = [-net]
    yr1_r = None
    for yr in range(1, ANALYSIS_YEARS + 1):
        b_yr = max(0.0, bat_kwh * (1 - BAT_DEG) ** yr)
        r    = simulate_grid_arbitrage(raw_df, b_yr, inv_kw, yr_offset=yr)
        if yr == 1:
            yr1_r = r
        base_yr = base * (1 + TARIFF_ESC) ** yr
        sav  = base_yr - r["net_cost"]
        savings.append(sav); net_costs.append(r["net_cost"])
        cum.append(cum[-1] + sav)
        cum_disc.append(cum_disc[-1] + sav / (1+OPPORTUNITY_RATE)**yr)

    pb_yr      = 0 if net == 0 else next((yr for yr, cf in enumerate(cum[1:], 1) if cf >= 0), None)
    pb_yr_disc = 0 if net == 0 else next((yr for yr, cf in enumerate(cum_disc[1:], 1) if cf >= 0), None)
    npv        = cum_disc[-1]
    return dict(label=label, solar_kw=0.0, bat_kwh=bat_kwh, inv_kw=inv_kw,
                cost=cost, fed=fed, state=state, net=net, tariff="Midday Saver",
                pb_yr=pb_yr, pb_yr_disc=pb_yr_disc, npv=npv,
                savings=savings, net_costs=net_costs, cum=cum, cum_disc=cum_disc,
                total_save=sum(savings),
                roi=sum(savings) / net * 100 if net > 0 else 0,
                yr1_save=savings[0] if savings else 0,
                self_suf_pct=yr1_r["self_suf_pct"] if yr1_r else 0.0,
                annual_export_kwh=0.0,
                annual_solar_kwh=0.0)


# ─── SOLAR SWEEP ─────────────────────────────────────────────────────────────
def sweep_solar_for_battery(raw_df, bat_kwh=14.0, tariff="Midday Saver",
                             solar_min=3.0, solar_max=20.0, step=0.5):
    """Step through solar sizes and record key metrics for a fixed battery size."""
    usable = bat_kwh * BAT_DOD
    base   = baseline(raw_df, tariff)
    rows   = []
    sizes  = np.arange(solar_min, solar_max + step / 2, step)
    for i, solar_kw in enumerate(sizes):
        sys.stdout.write(f"\r  Sweeping … {i+1}/{len(sizes)} ({solar_kw:.1f} kW)   ")
        sys.stdout.flush()
        inv_kw = min(solar_kw, 10.0)
        df     = add_solar(raw_df, solar_kw)
        r      = simulate(df, solar_kw, bat_kwh, inv_kw, tariff)
        daily  = r["df"].groupby("date").agg(bat_dis=("b_dis","sum"), bat_chg=("b_chg","sum"))
        avg_dis   = (daily["bat_dis"] * BAT_RTE**0.5).mean()
        bat_util  = min(avg_dis / usable * 100, 100.0)
        days_full = ((daily["bat_chg"] * BAT_RTE**0.5) >= usable * 0.90).mean() * 100
        rows.append(dict(
            solar_kw      = round(solar_kw, 1),
            inv_kw        = inv_kw,
            self_suf_pct  = r["self_suf_pct"],
            bat_util_pct  = bat_util,
            days_full_pct = days_full,
            export_kwh    = r["annual_export_kwh"],
            import_kwh    = r["annual_import_kwh"],
            net_cost      = r["net_cost"],
            annual_savings= base - r["net_cost"],
        ))
    print()
    return pd.DataFrame(rows)


def optimal_solar_size(sweep_df):
    """Knee of the self-sufficiency curve: first point where marginal gain drops below 2%/kW."""
    ss   = sweep_df["self_suf_pct"].values
    step = sweep_df["solar_kw"].iloc[1] - sweep_df["solar_kw"].iloc[0]
    grad = np.diff(ss) / step
    for i, g in enumerate(grad):
        if g < 2.0:
            return float(sweep_df["solar_kw"].iloc[i])
    return float(sweep_df.loc[sweep_df["self_suf_pct"].idxmax(), "solar_kw"])


def print_solar_sweep(sweep_df, bat_kwh, opt_kw, tariff):
    B = "="*78; S = "-"*78
    print(f"\n{B}")
    print(f"  SOLAR SWEEP — {bat_kwh} kWh battery | {tariff} | inverter capped at 10 kW")
    print(B)
    print(f"  {'Solar':>7}  {'Inv':>5}  {'Self-suf':>9}  {'Bat util':>9}  "
          f"{'Days full':>10}  {'Export':>9}  {'Saving/yr':>10}")
    print(f"  {S}")
    for _, row in sweep_df.iterrows():
        marker = "  ◄ recommended" if row["solar_kw"] == opt_kw else ""
        print(f"  {row['solar_kw']:>5.1f}kW  {row['inv_kw']:>4.0f}kW"
              f"  {row['self_suf_pct']:>8.1f}%"
              f"  {row['bat_util_pct']:>8.1f}%"
              f"  {row['days_full_pct']:>9.1f}%"
              f"  {row['export_kwh']:>8,.0f}kWh"
              f"  ${row['annual_savings']:>9,.0f}"
              f"{marker}")
    print(f"\n  → Recommended: {opt_kw:.1f} kW solar for a {bat_kwh} kWh battery")
    print(f"    (knee of self-sufficiency curve — diminishing returns beyond this point)\n")


def find_zero_net_combinations(raw_df, tariff="Midday Saver"):
    """For each battery size, find the minimum solar (kW) that achieves net_cost ≤ $0/yr.

    Sweeps solar 0–35 kW in 0.5 kW steps for each of the standard battery sizes.
    Returns a list of dicts with bat_kwh, min_solar_kw (None if not achievable ≤35 kW),
    net_cost, annual_export_kwh, and self_suf_pct.
    """
    # Zero-net = annual DEBS export credits >= import + supply charges
    # We sweep solar sizes until net_cost first hits $0 or below
    battery_sizes = [0.0, 5.0, 10.0, 13.5, 20.0, 30.0]
    solar_sizes   = np.arange(0.0, 35.5, 0.5)
    results       = []
    total_runs    = len(battery_sizes) * len(solar_sizes)
    run           = 0
    print(f"\n  Zero-net optimisation ({tariff}) …")
    for bat_kwh in battery_sizes:
        found = None
        for solar_kw in solar_sizes:
            run += 1
            sys.stdout.write(f"\r  [{run}/{total_runs}] {bat_kwh:.0f}kWh bat / {solar_kw:.1f}kW sol   ")
            sys.stdout.flush()
            inv_kw = min(max(solar_kw, 0.1), 10.0)
            df     = add_solar(raw_df, solar_kw)
            r      = simulate(df, solar_kw, bat_kwh, inv_kw, tariff)
            if r["net_cost"] <= 0 and found is None:
                found = dict(bat_kwh=bat_kwh, min_solar_kw=float(solar_kw),
                             net_cost=r["net_cost"],
                             annual_export_kwh=r["annual_export_kwh"],
                             self_suf_pct=r["self_suf_pct"])
                break   # found minimum for this battery size
        if found is None:
            # Record best achieved with max solar
            df  = add_solar(raw_df, solar_sizes[-1])
            r   = simulate(df, float(solar_sizes[-1]), bat_kwh, 10.0, tariff)
            found = dict(bat_kwh=bat_kwh, min_solar_kw=None,
                         net_cost=r["net_cost"],
                         annual_export_kwh=r["annual_export_kwh"],
                         self_suf_pct=r["self_suf_pct"])
        results.append(found)
    print()
    return results


def analyse_sizing(pbs, opt_kw=None):
    """Rank non-base solar+battery quotes by NPV efficiency and self-consumption.

    Given poor FiT rates (2–10¢/kWh vs ~32¢ import), exported kWh are worth far
    less than self-consumed kWh.  This function identifies the optimal sizing from
    the quoted options and quantifies the 'FiT penalty' from oversized arrays.

    Returns a dict with:
      rows        — list of dicts per Midday Saver solar quote, sorted by NPV desc
      best_npv    — quote with highest absolute NPV
      best_eff    — quote with highest NPV per $1,000 invested (capital efficiency)
      fit_penalty — avg value lost per exported kWh vs self-consuming at import rate
      opt_kw      — solar knee from sweep (if provided)
    """
    # Use Midday Saver only (optimal tariff with solar); exclude base case
    solar_pbs = [pb for pb in pbs
                 if pb["tariff"] == "Midday Saver"
                 and pb.get("solar_kw", 0) > 0
                 and pb["net"] > 0]
    if not solar_pbs:
        return None

    # Effective FiT rate: weighted avg of DEBS peak (10¢) and off-peak (2¢) export credits.
    # Peak window = 6 hrs/day out of 24 → roughly 25% of exports at peak, 75% at off-peak.
    avg_fit_rate = 0.25 * DEBS_PEAK + 0.75 * DEBS_OP      # ≈ 4.0¢/kWh
    fit_penalty  = A1_RATE - avg_fit_rate                   # value lost per exported kWh

    rows = []
    for pb in solar_pbs:
        solar_kwh  = pb.get("annual_solar_kwh", 0.0)
        export_kwh = pb.get("annual_export_kwh", 0.0)
        export_ratio = export_kwh / solar_kwh * 100 if solar_kwh > 0 else 0.0
        # Annual value destroyed by exporting instead of self-consuming
        annual_fit_loss = export_kwh * fit_penalty
        npv_per_k = pb["npv"] / (pb["net"] / 1000) if pb["net"] > 0 else 0.0
        rows.append(dict(
            label         = pb["label"],
            solar_kw      = pb["solar_kw"],
            bat_kwh       = pb["bat_kwh"],
            net           = pb["net"],
            npv           = pb["npv"],
            self_suf_pct  = pb.get("self_suf_pct", 0.0),
            export_kwh    = export_kwh,
            solar_kwh     = solar_kwh,
            export_ratio  = export_ratio,
            annual_fit_loss = annual_fit_loss,
            npv_per_k     = npv_per_k,
        ))

    rows.sort(key=lambda x: -x["npv"])
    best_npv  = max(rows, key=lambda x: x["npv"])  if rows else None
    best_eff  = max(rows, key=lambda x: x["npv_per_k"]) if rows else None
    return dict(rows=rows, best_npv=best_npv, best_eff=best_eff,
                fit_penalty=fit_penalty, avg_fit_rate=avg_fit_rate, opt_kw=opt_kw)


def print_zero_net(results, tariff):
    B = "="*78; S = "-"*78
    print(f"\n{B}")
    print(f"  ZERO-NET SIZING — minimum solar for ≤$0 net electricity cost/yr ({tariff})")
    print(B)
    print(f"  {'Battery':>10}  {'Min solar':>10}  {'Net cost':>11}  {'Export':>9}  {'Self-suf':>9}")
    print(f"  {S}")
    for r in results:
        solar_str = f"{r['min_solar_kw']:.1f} kW" if r['min_solar_kw'] is not None else ">35 kW *"
        note      = "" if r['min_solar_kw'] is not None else "  (not achievable)"
        print(f"  {r['bat_kwh']:>8.1f} kWh"
              f"  {solar_str:>10}"
              f"  ${r['net_cost']:>10,.0f}"
              f"  {r['annual_export_kwh']:>8,.0f} kWh"
              f"  {r['self_suf_pct']:>8.1f}%"
              f"{note}")
    print(f"\n  * Net cost shown is for 35 kW solar — zero-net not achievable at that battery size.\n"
          f"    Note: feed-in tariff offset (DEBS) counts toward 'zero net cost' but DEBS rates\n"
          f"    decline {DEBS_DECL*100:.0f}% per year. The model uses year-0 rates here.\n")


# ─── COLOURS ──────────────────────────────────────────────────────────────────
CS = dict(solar="#F5A623", bdis="#7ED321", bchg="#4A90D9",
          gimp="#D0021B", gexp="#9013FE")
QC = ["#e8463a","#0f9d58","#f5a623","#9c27b0","#00bcd4"]
CA1="#1a73e8"; CMS="#0f9d58"


# ─── MONTHLY NET HELPER ───────────────────────────────────────────────────────
def _monthly_net(res):
    """Return a 12-row DataFrame (index 1–12) with monthly net electricity cost.

    Columns: ic (import cost), ec (export credit), days, net (= ic + days*supply - ec).
    Works for any simulate() or simulate_grid_arbitrage() result dict.
    """
    df = res["df"].copy()
    df["month"] = df["datetime"].dt.month
    df["date"]  = df["datetime"].dt.date

    slots = df["slot"].values.astype(int)

    if res["tariff"] == "A1 Flat":
        df["ic"] = df["g_imp"] * A1_RATE
        sc = A1_SUPPLY
    else:
        # vectorised Midday Saver rate lookup — avoids slow per-row apply()
        df["ic"] = df["g_imp"] * np.vectorize(ms_rate)(slots)
        sc = MS_SUPPLY

    df["ec"] = df["g_exp"] * np.vectorize(debs_rate)(slots)

    m = df.groupby("month").agg(ic=("ic","sum"), ec=("ec","sum"),
                                days=("date","nunique"))
    m = m.reindex(range(1, 13), fill_value=0)
    m["net"] = m["ic"] + m["days"] * sc - m["ec"]
    m["sc"] = sc   # supply charge $/day, handy for callers
    return m


# ─── PLOTS ────────────────────────────────────────────────────────────────────
def avg_day_plot(res, ax, title_suffix=""):
    df  = res["df"]
    avg = df.groupby("slot")[["solar_kwh","s_slf","b_dis","g_imp","consumption_kwh"]].mean()
    h   = avg.index/2
    ax.stackplot(h, avg["s_slf"], avg["b_dis"], avg["g_imp"],
                 labels=["Solar self-consumed","Battery discharge","Grid import"],
                 colors=[CS["solar"],CS["bdis"],CS["gimp"]], alpha=0.85)
    ax.plot(h, avg["consumption_kwh"], "k-", lw=1.3, label="Consumption")
    ax.plot(h, avg["solar_kwh"], "--", color=CS["solar"], lw=0.9, alpha=0.7,
            label="Solar gen (shaded)")
    ax.axvspan(15,21,alpha=0.07,color="red",  label="Peak 3–9pm")
    ax.axvspan(9, 15,alpha=0.06,color="green",label="Super off-peak")
    ax.set_xlim(0,24); ax.set_xticks(range(0,25,3))
    ax.set_xticklabels([f"{h:02d}:00" for h in range(0,25,3)],fontsize=8)
    ax.set_xlabel("Hour"); ax.set_ylabel("Avg kWh/30min")
    ax.set_title(f"{res['solar_kw']}kW solar / {res['bat_kwh']}kWh battery "
                 f"({res['tariff']}) {title_suffix}", fontsize=9)
    ax.legend(loc="upper left", fontsize=7, ncol=3)


def seasonal_plots(res, axes):
    df = res["df"].copy()
    df["month"] = df["datetime"].dt.month
    seasons = [("Summer (Dec–Feb)",[12,1,2]),("Autumn (Mar–May)",[3,4,5]),
               ("Winter (Jun–Aug)",[6,7,8]),("Spring (Sep–Nov)",[9,10,11])]
    for ax,(lbl,months) in zip(axes.flat, seasons):
        sub = df[df["month"].isin(months)]
        if sub.empty: continue
        avg = sub.groupby("slot")[["solar_kwh","s_slf","b_dis","g_imp","consumption_kwh"]].mean()
        h   = avg.index/2
        ax.stackplot(h, avg["s_slf"],avg["b_dis"],avg["g_imp"],
                     colors=[CS["solar"],CS["bdis"],CS["gimp"]], alpha=0.85)
        ax.plot(h, avg["consumption_kwh"],"k-",lw=1.1)
        ax.plot(h, avg["solar_kwh"],"--",color=CS["solar"],lw=0.8,alpha=0.7)
        ax.axvspan(15,21,alpha=0.07,color="red")
        ax.axvspan(9,15,alpha=0.06,color="green")
        ax.set_xlim(0,24); ax.set_xticks(range(0,25,6))
        ax.set_xticklabels([f"{h:02d}h" for h in range(0,25,6)],fontsize=8)
        ax.set_title(lbl, fontsize=9); ax.set_ylabel("Avg kWh/30min",fontsize=8)


def weekly_plot(res, week_start, axes):
    df   = res["df"]
    mask = (df["datetime"]>=week_start)&(df["datetime"]<week_start+pd.Timedelta(days=7))
    w    = df[mask]
    if w.empty: return
    x    = w["datetime"]
    ax   = axes[0]
    ax.stackplot(x, w["s_slf"],w["b_dis"],w["g_imp"],
                 labels=["Solar self","Battery dis","Grid import"],
                 colors=[CS["solar"],CS["bdis"],CS["gimp"]], alpha=0.85)
    ax.plot(x, w["consumption_kwh"],"k-",lw=0.8,label="Consumption")
    ax.plot(x, w["solar_kwh"],"--",color=CS["solar"],lw=0.7,alpha=0.7,label="Solar gen")
    ax.set_ylabel("kWh/30min")
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%a %d %b"))
    ax.xaxis.set_major_locator(mdates.DayLocator())
    plt.setp(ax.xaxis.get_majorticklabels(), rotation=30, ha="right")
    ax.legend(loc="upper right",fontsize=7,ncol=5)
    ax.set_title(f"Week {week_start.strftime('%d %b')} — {res['solar_kw']}kW/{res['bat_kwh']}kWh ({res['tariff']})",fontsize=9)
    ax2  = axes[1]
    ax2.fill_between(x, w["soc"], alpha=0.6, color=CS["bchg"])
    ax2.axhline(res["bat_kwh"]*BAT_DOD,color="grey",lw=0.7,ls="--",label="Max DoD")
    ax2.set_ylabel("Battery SoC (kWh)"); ax2.set_ylim(0, res["bat_kwh"]+0.5)
    ax2.xaxis.set_major_formatter(mdates.DateFormatter("%a"))
    ax2.xaxis.set_major_locator(mdates.DayLocator())
    plt.setp(ax2.xaxis.get_majorticklabels(), rotation=30, ha="right")
    ax2.legend(fontsize=7)


def monthly_cost_plot(res_a1, res_ms, raw_df, ax):
    lbls = ["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"]
    x    = np.arange(12)
    for idx, res in enumerate([res_a1, res_ms]):
        m = _monthly_net(res)
        ax.bar(x+idx*0.35, m["net"].values, 0.35,
               label=res["tariff"], color=[CA1,CMS][idx], alpha=0.85)
    # Baseline: no solar/battery — use raw consumption as grid import
    base_res = {"df": raw_df.assign(g_imp=raw_df["consumption_kwh"], g_exp=0.0),
                "tariff": "A1 Flat"}
    mb = _monthly_net(base_res)
    ax.step(x+0.35, mb["net"].values, where="mid", color="black", lw=1.5, ls="--",
            label="Baseline (grid only / A1)")
    ax.set_xticks(x+0.175); ax.set_xticklabels(lbls)
    ax.set_ylabel("Monthly Net Cost (AUD)")
    ax.set_title("Monthly Cost — Baseline vs Solar+Battery")
    ax.legend(fontsize=8)
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda v,_: f"${v:.0f}"))


def monthly_cost_all_plot(all_res, ax_a1, ax_ms):
    """Monthly cost for every quote on both tariffs — base case highlighted.

    Use this to verify the modelled base-case bill against your actual statements.
    Base case = thick solid line with markers; solar+battery quotes = thinner dashed lines.
    """
    month_lbls = ["Jan","Feb","Mar","Apr","May","Jun",
                  "Jul","Aug","Sep","Oct","Nov","Dec"]
    x      = np.arange(12)

    for ax, tariff in [
        (ax_a1, "A1 Flat"),
        (ax_ms, "Midday Saver"),
    ]:
        quote_results = [r for r in all_res if r["tariff"] == tariff]
        for idx, r in enumerate(quote_results):
            m = _monthly_net(r)

            is_base = r["solar_kw"] == 0 and r.get("bat_kwh", 0) == 0
            col     = QC[idx % len(QC)]
            ax.plot(x, m["net"].values,
                    color=col,
                    lw=2.5 if is_base else 1.4,
                    ls="-"  if is_base else "--",
                    marker="o" if is_base else None,
                    markersize=5,
                    zorder=5 if is_base else 2,
                    label=r["label"] + (" ← verify against your bills" if is_base else ""))

        ax.set_xticks(x)
        ax.set_xticklabels(month_lbls)
        ax.set_ylabel("Monthly Net Cost (AUD)")
        ax.set_ylim(bottom=0)
        ax.set_title(f"{tariff} — Monthly Cost: Base Case vs Solar+Battery\n"
                     f"Base case (solid + dots) should match your actual bills")
        ax.legend(fontsize=7)
        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda v, _: f"${v:,.0f}"))
        ax.grid(axis="y", alpha=0.3)


def payback_plot(pbs, ax):
    """Cumulative cashflow for each solar+battery quote over the analysis horizon."""
    yrs = list(range(ANALYSIS_YEARS + 1))

    for idx, pb in enumerate(pbs):
        col = QC[idx % len(QC)]
        ls  = "-" if pb["tariff"] == "A1 Flat" else "--"

        lbl = f"{pb['label']} | {pb['tariff']}"
        if pb["pb_yr"]:
            lbl += f"  ← payback yr {pb['pb_yr']}"
        ax.plot(yrs, pb["cum"], color=col, ls=ls, lw=1.8, label=lbl)

    ax.axhline(0, color="black", lw=0.8, ls=":")
    ax.fill_between(yrs, 0, color="green", alpha=0.04)
    ax.set_xlabel("Year")
    ax.set_ylabel("Cumulative Cash-Flow (AUD)")
    ax.set_title(
        f"Solar+Battery Cumulative Cash-Flow — {ANALYSIS_YEARS}-Year Horizon\n"
        f"Solid = A1 Flat, Dashed = Midday Saver")
    ax.legend(fontsize=7)
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda v, _: f"${v:,.0f}"))
    ax.xaxis.set_major_locator(plt.MultipleLocator(5))
    ax.grid(axis="y", alpha=0.3)


def tariff_comp_plot(ra1, rms, ax):
    lbls = ["Import+Supply","Export Credit","Net Annual Cost"]
    a1v  = [ra1["import_cost"]+ra1["supply_cost"], ra1["export_credit"], ra1["net_cost"]]
    msv  = [rms["import_cost"]+rms["supply_cost"], rms["export_credit"], rms["net_cost"]]
    x    = np.arange(3); w=0.35
    b1   = ax.bar(x-w/2, a1v, w, label="A1 Flat",     color=CA1, alpha=0.85)
    b2   = ax.bar(x+w/2, msv, w, label="Midday Saver", color=CMS, alpha=0.85)
    for b in list(b1)+list(b2):
        ax.text(b.get_x()+b.get_width()/2, b.get_height()+15,
                f"${b.get_height():,.0f}", ha="center", va="bottom", fontsize=8)
    ax.set_xticks(x); ax.set_xticklabels(lbls)
    ax.set_ylabel("AUD")
    ax.set_title(f"Tariff Comparison — {ra1['solar_kw']}kW / {ra1['bat_kwh']}kWh")
    ax.legend()
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda v,_: f"${v:,.0f}"))


def heatmap_plot(res, ax, label):
    df = res["df"].copy()
    df["month"] = df["datetime"].dt.month
    df["hour"]  = df["datetime"].dt.hour
    piv = df.groupby(["hour","month"])["g_imp"].mean().unstack(fill_value=0)
    im  = ax.imshow(piv.values, aspect="auto", cmap="YlOrRd", origin="lower")
    ax.set_xticks(range(12))
    ax.set_xticklabels(["J","F","M","A","M","J","J","A","S","O","N","D"],fontsize=8)
    ax.set_yticks(range(0,24,3))
    ax.set_yticklabels([f"{h:02d}:00" for h in range(0,24,3)],fontsize=8)
    ax.set_xlabel("Month"); ax.set_ylabel("Hour")
    ax.set_title(f"Grid Import Heat-Map — {label}\n{res['solar_kw']}kW/{res['bat_kwh']}kWh",fontsize=9)
    plt.colorbar(im, ax=ax, label="kWh")
    ax.axhspan(14.5,20.5,color="blue",alpha=0.07)


def shading_plot(ax):
    slots = np.arange(48); hours = slots/2
    for (lbl,doy),col in zip([("Summer Jan",1),("Autumn Apr",92),
                               ("Winter Jul",183),("Spring Oct",275)],
                              ["#F5A623","#7ED321","#4A90D9","#9013FE"]):
        facs = np.array([shading_factor(int(s),doy) for s in slots])*100
        ax.plot(hours, facs, label=lbl, color=col, lw=1.8)
    ax.axvspan(15,21,alpha=0.07,color="red",  label="Peak 3–9pm")
    ax.axvspan(9,15, alpha=0.06,color="green",label="Super off-peak")
    ax.axhline(40,color="black",lw=0.8,ls=":",label="Stated avg ~40%")
    ax.set_xlim(0,24); ax.set_xticks(range(0,25,3))
    ax.set_xticklabels([f"{h:02d}:00" for h in range(0,25,3)],fontsize=8)
    ax.set_ylabel("% of unshaded output"); ax.set_xlabel("Hour of day")
    ax.set_title("Site Shading Model — Seasonal & Diurnal\n(deeper in winter, tapering toward sunrise/sunset)")
    ax.legend(fontsize=8, ncol=3); ax.set_ylim(0,100); ax.grid(axis="y",alpha=0.3)


def monthly_savings_plot(all_res, raw_df, ax):
    """Bar chart: monthly saving vs no-solar baseline for every quote/tariff combo."""
    month_lbls = ["Jan","Feb","Mar","Apr","May","Jun",
                  "Jul","Aug","Sep","Oct","Nov","Dec"]

    # Monthly baseline for each tariff — build a no-solar result dict so _monthly_net can compute it
    rd = raw_df.copy()
    rd["month"] = rd["datetime"].dt.month
    rd["date"]  = rd["datetime"].dt.date
    rd["ic_a1"] = rd["consumption_kwh"] * A1_RATE
    rd["ic_ms"] = rd["consumption_kwh"] * np.vectorize(ms_rate)(rd["slot"].values.astype(int))
    mb = rd.groupby("month").agg(
        ic_a1=("ic_a1","sum"), ic_ms=("ic_ms","sum"), days=("date","nunique"))
    mb["base_a1"] = mb["ic_a1"] + mb["days"] * A1_SUPPLY
    mb["base_ms"] = mb["ic_ms"] + mb["days"] * MS_SUPPLY
    mb = mb.reindex(range(1, 13), fill_value=0)

    x     = np.arange(12)
    n     = len(all_res)
    width = 0.8 / n

    for idx, r in enumerate(all_res):
        m       = _monthly_net(r)
        base_m  = mb["base_a1"] if r["tariff"] == "A1 Flat" else mb["base_ms"]
        savings = base_m.values - m["net"].values
        col         = QC[idx // 2 % len(QC)]
        hatch       = "" if r["tariff"] == "A1 Flat" else "///"
        label       = f"{r['label'][:24]} | {r['tariff']}"
        offset      = (idx - n / 2 + 0.5) * width
        ax.bar(x + offset, savings, width, label=label,
               color=col, hatch=hatch, alpha=0.85, edgecolor="white")

    ax.axhline(0, color="black", lw=0.8, ls=":")
    ax.set_xticks(x)
    ax.set_xticklabels(month_lbls)
    ax.set_ylabel("Monthly Saving vs Baseline (AUD)")
    ax.set_title("Expected Monthly Savings — All Quotes & Tariffs\n"
                 "Solid = A1 Flat  |  Hatched = Midday Saver")
    ax.legend(fontsize=7, ncol=2)
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda v, _: f"${v:,.0f}"))
    ax.grid(axis="y", alpha=0.3)


def monthly_bill_comparison_plot(all_res, arbit_res, raw_df, ax):
    """Line chart of predicted monthly net cost for every scenario on one axes.

    Baseline (grid-only A1 Flat) is plotted as a thick black reference line.
    Solar+battery quotes are coloured lines; solid = A1 Flat, dashed = Midday Saver.
    Battery arbitrage (no solar) is shown in grey.
    """
    month_lbls = ["Jan","Feb","Mar","Apr","May","Jun",
                  "Jul","Aug","Sep","Oct","Nov","Dec"]
    x      = np.arange(12)

    # ── Baseline: compute monthly cost as if there's no solar/battery on A1 Flat
    # We fake a minimal result dict with the raw consumption data so _monthly_net() can handle it
    base_res = {"df": raw_df.assign(g_imp=raw_df["consumption_kwh"], g_exp=0.0),
                "tariff": "A1 Flat"}
    mb = _monthly_net(base_res)
    ax.plot(x, mb["net"].values, "k-", lw=2.5, marker="o", ms=5,
            zorder=5, label="Baseline — grid only (A1 Flat)")

    # ── All solar+battery quotes + arbitrage
    colour_idx = 0
    all_scenarios = list(all_res) + (list(arbit_res) if arbit_res else [])
    for r in all_scenarios:
        is_base = r["solar_kw"] == 0 and r.get("bat_kwh", 0) == 0
        if is_base:
            continue

        m = _monthly_net(r)

        is_arb = r["solar_kw"] == 0 and r.get("bat_kwh", 0) > 0
        if is_arb:
            col = "#7f8c8d"     # grey for battery-only arbitrage
        else:
            col = QC[colour_idx % len(QC)]
            colour_idx += 1

        ls  = "-" if r["tariff"] == "A1 Flat" else "--"
        lbl = f"{r['label'][:28]} | {r['tariff']}"
        ax.plot(x, m["net"].values, color=col, ls=ls, lw=1.8, label=lbl)

    ax.set_xticks(x)
    ax.set_xticklabels(month_lbls)
    ax.set_ylabel("Monthly Net Cost (AUD)")
    ax.set_ylim(bottom=0)
    ax.set_title("Predicted Monthly Electricity Bill — All Scenarios\n"
                 "Black = current baseline  ·  Solid = A1 Flat  ·  Dashed = Midday Saver")
    ax.legend(fontsize=7, ncol=2, loc="upper right")
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda v, _: f"${v:,.0f}"))
    ax.grid(axis="y", alpha=0.3)


def arbitrage_day_plot(res, ax1, ax2):
    """Average daily profile for a battery-only grid arbitrage scenario.

    Battery charges from grid during SOP (green band) and discharges at peak (red band).
    The top axes decompose grid import into 'load only' vs 'battery charging', showing
    the energy shift that makes arbitrage economical.
    """
    df  = res["df"]
    avg = df.groupby("slot")[["b_chg", "b_dis", "g_imp", "consumption_kwh", "soc"]].mean()
    h   = avg.index / 2.0
    rte = BAT_RTE ** 0.5

    load_from_bat  = avg["b_dis"] * rte                           # energy delivered to load
    load_from_grid = (avg["g_imp"] - avg["b_chg"]).clip(lower=0)  # grid covering load only
    bat_chg_draw   = avg["b_chg"]                                  # extra grid draw for charging

    ax1.stackplot(h, load_from_bat, load_from_grid,
                  labels=["Battery discharge (to load)", "Grid import (for load)"],
                  colors=[CS["bdis"], CS["gimp"]], alpha=0.85)
    ax1.fill_between(h, 0, bat_chg_draw, alpha=0.45, color=CS["bchg"],
                     hatch="///", label="Grid draw (to charge battery, SOP)")
    ax1.plot(h, avg["consumption_kwh"], "k-", lw=1.3, label="Consumption")
    ax1.axvspan(9, 15,  alpha=0.08, color="green", label="SOP 9am–3pm (8.6¢/kWh — charge)")
    ax1.axvspan(15, 21, alpha=0.07, color="red",   label="Peak 3–9pm (53.8¢/kWh — discharge)")
    ax1.set_xlim(0, 24)
    ax1.set_xticks(range(0, 25, 3))
    ax1.set_xticklabels([f"{hh:02d}:00" for hh in range(0, 25, 3)], fontsize=8)
    ax1.set_ylabel("Avg kWh / 30 min")
    ax1.set_title(f"Battery Grid Arbitrage — {res['bat_kwh']} kWh Battery, No Solar\n"
                  "Charges cheap during SOP (green), discharges at peak (red)", fontsize=9)
    ax1.legend(loc="upper left", fontsize=7, ncol=2)

    ax2.fill_between(h, avg["soc"], alpha=0.6, color=CS["bchg"], label="Battery SoC (avg)")
    b_cap = res["bat_kwh"] * BAT_DOD
    ax2.axhline(b_cap, color="grey", lw=0.9, ls="--", label=f"Usable capacity ({b_cap:.1f} kWh)")
    ax2.axvspan(9, 15,  alpha=0.07, color="green")
    ax2.axvspan(15, 21, alpha=0.06, color="red")
    ax2.set_xlim(0, 24)
    ax2.set_xticks(range(0, 25, 3))
    ax2.set_xticklabels([f"{hh:02d}:00" for hh in range(0, 25, 3)], fontsize=8)
    ax2.set_ylabel("Battery SoC (kWh)")
    ax2.set_ylim(0, res["bat_kwh"] + 0.5)
    ax2.set_xlabel("Hour of day")
    ax2.legend(fontsize=7)
    ax2.grid(axis="y", alpha=0.3)


def zero_net_plot(results, tariff, ax):
    """Bar chart: minimum solar kW needed per battery size for zero net electricity cost."""
    labels  = [f"{r['bat_kwh']:.0f} kWh" for r in results]
    values  = [r["min_solar_kw"] if r["min_solar_kw"] is not None else 36.0 for r in results]
    colors  = ["#2ecc71" if r["min_solar_kw"] is not None else "#e74c3c" for r in results]

    bars = ax.bar(labels, values, color=colors, alpha=0.85, edgecolor="white", width=0.5)
    for bar, r in zip(bars, results):
        if r["min_solar_kw"] is not None:
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.3,
                    f"{r['min_solar_kw']:.1f} kW", ha="center", va="bottom",
                    fontsize=9, fontweight="bold")
        else:
            ax.text(bar.get_x() + bar.get_width() / 2, 1.0,
                    ">35 kW\n(not achievable)", ha="center", va="bottom",
                    fontsize=8, color="white", fontweight="bold")

    ax.set_xlabel("Battery size")
    ax.set_ylabel("Minimum solar array (kW)")
    ax.set_ylim(0, max(max(values), 5) + 3)
    ax.set_title(f"Zero-Net Sizing — Minimum Solar for $0 Annual Electricity Cost\n"
                 f"Tariff: {tariff}  ·  Green = achievable ≤35 kW  ·  Red = not achievable within model")
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda v, _: f"{v:.0f} kW"))
    ax.grid(axis="y", alpha=0.3)
    handles = [plt.Rectangle((0, 0), 1, 1, color="#2ecc71", alpha=0.85, label="Zero-net achievable"),
               plt.Rectangle((0, 0), 1, 1, color="#e74c3c", alpha=0.85, label=">35 kW needed")]
    ax.legend(handles=handles, fontsize=8)


def consumption_heatmap(raw_df, ax):
    df = raw_df.copy()
    df["month"] = df["datetime"].dt.month
    df["hour"]  = df["datetime"].dt.hour
    piv = df.groupby(["hour","month"])["consumption_kwh"].mean().unstack(fill_value=0)
    im  = ax.imshow(piv.values, aspect="auto", cmap="Blues", origin="lower")
    ax.set_xticks(range(12))
    ax.set_xticklabels(["J","F","M","A","M","J","J","A","S","O","N","D"],fontsize=8)
    ax.set_yticks(range(0,24,3))
    ax.set_yticklabels([f"{h:02d}:00" for h in range(0,24,3)],fontsize=8)
    ax.set_xlabel("Month"); ax.set_ylabel("Hour")
    ax.set_title("Your Actual Consumption — Heat-Map (kWh/30min)\n(red band = peak tariff window 3–9pm)")
    plt.colorbar(im, ax=ax, label="kWh")
    ax.axhspan(14.5,20.5,color="red",alpha=0.07)


def solar_sweep_plot(sweep_df, bat_kwh, tariff, opt_kw, ax):
    x    = sweep_df["solar_kw"]
    col  = dict(ss="#F5A623", bat="#0f9d58", exp="#D0021B")
    l1,  = ax.plot(x, sweep_df["self_suf_pct"],  color=col["ss"],  lw=2,   label="Self-sufficiency %")
    l2,  = ax.plot(x, sweep_df["bat_util_pct"],  color=col["bat"], lw=2,   label="Avg daily battery utilisation %")
    l3,  = ax.plot(x, sweep_df["days_full_pct"], color=col["bat"], lw=1.5, ls="--", label="Days battery fully charged %")
    ax.set_xlabel("Solar array size (kW)")
    ax.set_ylabel("Percentage (%)")
    ax.set_ylim(0, 105)
    ax.grid(axis="y", alpha=0.3)
    ax2   = ax.twinx()
    l4,  = ax2.plot(x, sweep_df["export_kwh"], color=col["exp"], lw=1.5, ls=":", label="Annual export (kWh)")
    ax2.set_ylabel("Annual export (kWh)", color=col["exp"])
    ax2.tick_params(axis="y", labelcolor=col["exp"])
    ax.axvline(opt_kw, color="black", lw=1.2, ls="--", alpha=0.7)
    ax.text(opt_kw + 0.2, 8, f"← recommended\n   {opt_kw:.1f} kW", fontsize=8, va="bottom")
    ax.legend([l1, l2, l3, l4], [l.get_label() for l in [l1, l2, l3, l4]],
              fontsize=8, loc="lower right")
    ax.set_title(f"Solar Size Optimisation — {bat_kwh} kWh Battery ({tariff})\n"
                 f"Inverter sized to match solar (max 10 kW); self-sufficiency knee marks sweet spot")


def seasonal_day_plot(res, months, ax1, ax2):
    """Plot energy flows (ax1) and battery SoC (ax2) for a representative season day.

    Picks the day whose total consumption is closest to the seasonal median —
    avoids anomalous high/low days while still using real data.
    Returns the chosen date, or None if no data exists for those months.
    """
    df = res["df"].copy()
    df["month"] = df["datetime"].dt.month
    sub = df[df["month"].isin(months)]
    if sub.empty:
        return None

    daily_kwh = sub.groupby("date")["consumption_kwh"].sum()
    date      = (daily_kwh - daily_kwh.median()).abs().idxmin()
    day       = df[df["date"] == date].sort_values("slot")
    hours     = day["slot"] / 2.0

    # Energy flows — stackplot fills the load shape
    ax1.stackplot(hours,
                  day["s_slf"],
                  day["b_dis"] * BAT_RTE**0.5,
                  day["g_imp"],
                  labels=["Solar self-consumed", "Battery discharge", "Grid import"],
                  colors=[CS["solar"], CS["bdis"], CS["gimp"]], alpha=0.85)
    ax1.plot(hours, day["consumption_kwh"], "k-",  lw=1.5, label="Consumption")
    ax1.plot(hours, day["solar_kwh"],       "--",  lw=1.0, color=CS["solar"], alpha=0.7,
             label="Solar gen (shaded)")
    ax1.fill_between(hours, day["g_exp"], alpha=0.35, color=CS["gexp"], label="Grid export")
    ax1.axvspan(15, 21, alpha=0.07, color="red",   label="Peak 3–9pm")
    ax1.axvspan(9,  15, alpha=0.06, color="green", label="Super off-peak")
    ax1.set_xlim(0, 24)
    ax1.set_xticks(range(0, 25, 3))
    ax1.set_xticklabels([f"{h:02d}:00" for h in range(0, 25, 3)], fontsize=8)
    ax1.set_ylabel("kWh / 30min")
    ax1.legend(loc="upper left", fontsize=7, ncol=4)
    ax1.grid(axis="y", alpha=0.3)

    # Battery SoC
    ax2.fill_between(hours, day["soc"], alpha=0.6, color=CS["bchg"], label="Battery SoC")
    ax2.axhline(res["bat_kwh"] * BAT_DOD, color="grey", lw=0.9, ls="--", label="Usable capacity")
    ax2.set_xlim(0, 24)
    ax2.set_xticks(range(0, 25, 3))
    ax2.set_xticklabels([f"{h:02d}:00" for h in range(0, 25, 3)], fontsize=8)
    ax2.set_ylabel("Battery SoC (kWh)")
    ax2.set_ylim(0, res["bat_kwh"] + 0.5)
    ax2.set_xlabel("Hour of day")
    ax2.legend(fontsize=7)
    ax2.grid(axis="y", alpha=0.3)

    return date


# ─── DASHBOARD ────────────────────────────────────────────────────────────────
def make_dashboard(all_res, pbs, raw_df, out_dir,
                   sweep_df=None, sweep_bat_kwh=14.0, sweep_tariff="Midday Saver", sweep_opt_kw=None,
                   arbit_res=None, arbit_pbs=None, zero_net=None, zero_net_tariff="Midday Saver"):
    out_dir.mkdir(parents=True, exist_ok=True)
    # Use first quote that actually has solar for profile/seasonal/heatmap charts
    ra1 = next(r for r in all_res if r["tariff"] == "A1 Flat"     and r["solar_kw"] > 0)
    rms = next(r for r in all_res if r["tariff"] == "Midday Saver" and r["solar_kw"] > 0)
    yr = raw_df["datetime"].dt.year.mode()[0]

    def save(fig, name):
        fig.savefig(out_dir/name, dpi=150, bbox_inches="tight")
        plt.close(fig)
        print(f"  Saved: {name}")

    # 00 consumption + shading
    fig,axes = plt.subplots(1,2,figsize=(15,6))
    fig.suptitle("Your Actual Consumption & Site Shading Model",fontsize=13,fontweight="bold")
    consumption_heatmap(raw_df, axes[0]); shading_plot(axes[1])
    plt.tight_layout(); save(fig,"00_consumption_and_shading.png")

    # 01 average daily — full year
    fig,axes = plt.subplots(2,1,figsize=(14,10))
    fig.suptitle("Average Daily Profile (Full Year) — Your Actual Data",fontsize=13,fontweight="bold")
    avg_day_plot(ra1,axes[0]); avg_day_plot(rms,axes[1])
    plt.tight_layout(); save(fig,"01_average_daily_profile.png")

    # 02 seasonal
    fig,axes = plt.subplots(2,2,figsize=(15,10))
    fig.suptitle(f"Seasonal Profiles — {ra1['solar_kw']}kW/{ra1['bat_kwh']}kWh (A1 Flat)",
                 fontsize=13,fontweight="bold")
    seasonal_plots(ra1, axes)
    handles = [plt.Rectangle((0,0),1,1,color=CS["solar"],alpha=0.85,label="Solar self"),
               plt.Rectangle((0,0),1,1,color=CS["bdis"], alpha=0.85,label="Battery dis"),
               plt.Rectangle((0,0),1,1,color=CS["gimp"], alpha=0.85,label="Grid import"),
               plt.Line2D([0],[0],color="black",lw=1.2,label="Consumption"),
               plt.Line2D([0],[0],color=CS["solar"],lw=0.9,ls="--",label="Solar gen (shaded)")]
    fig.legend(handles=handles,loc="lower center",ncol=5,fontsize=8,bbox_to_anchor=(0.5,0))
    plt.tight_layout(rect=[0,0.05,1,1]); save(fig,"02_seasonal_profiles.png")

    # 03 weekly summer + winter
    for slbl,(mo,dy) in [("Summer",(1,15)),("Winter",(7,15))]:
        ws  = pd.Timestamp(year=yr,month=mo,day=dy)
        fig,axes = plt.subplots(4,1,figsize=(15,14))
        fig.suptitle(f"{slbl} Week {ws.strftime('%d %b %Y')} — Your Actual Data",
                     fontsize=13,fontweight="bold")
        weekly_plot(ra1,ws,[axes[0],axes[1]])
        weekly_plot(rms,ws,[axes[2],axes[3]])
        plt.tight_layout(); save(fig,f"03_{slbl.lower()}_week.png")

    # 04 monthly costs (first solar quote — summary chart)
    fig,ax = plt.subplots(figsize=(13,6))
    fig.suptitle("Monthly Electricity Cost — Your Actual Data",fontsize=13,fontweight="bold")
    monthly_cost_plot(ra1,rms,raw_df,ax)
    plt.tight_layout(); save(fig,"04_monthly_costs.png")

    # 14 monthly cost — all quotes vs base case (verification chart)
    fig,axes = plt.subplots(1,2,figsize=(16,6),sharey=True)
    fig.suptitle("Monthly Cost — All Quotes vs Base Case\n"
                 "Base case (solid line + dots) should match your actual electricity bills",
                 fontsize=13,fontweight="bold")
    monthly_cost_all_plot(all_res,axes[0],axes[1])
    plt.tight_layout(); save(fig,"14_monthly_cost_verification.png")

    # 05 tariff comparison
    fig,ax = plt.subplots(figsize=(10,6))
    fig.suptitle("Annual Cost: A1 Flat vs Midday Saver",fontsize=13,fontweight="bold")
    tariff_comp_plot(ra1,rms,ax)
    plt.tight_layout(); save(fig,"05_tariff_comparison.png")

    # 06 payback — all quotes + arbitrage
    fig,ax = plt.subplots(figsize=(14,7))
    fig.suptitle("Payback Analysis — All Quotes (25-Year Horizon)",
                 fontsize=13,fontweight="bold")
    combined_pbs = list(pbs) + (list(arbit_pbs) if arbit_pbs else [])
    payback_plot(combined_pbs, ax)
    plt.tight_layout(); save(fig,"06_payback_analysis.png")

    # 07 heatmaps
    fig,axes = plt.subplots(1,2,figsize=(15,7))
    fig.suptitle("Grid Import Heat-Map — A1 vs Midday Saver",fontsize=13,fontweight="bold")
    heatmap_plot(ra1,axes[0],"A1 Flat"); heatmap_plot(rms,axes[1],"Midday Saver")
    plt.tight_layout(); save(fig,"07_import_heatmap.png")

    # 08 solar sweep
    if sweep_df is not None:
        fig,ax = plt.subplots(figsize=(13,6))
        fig.suptitle(f"Solar Size Optimisation — {sweep_bat_kwh} kWh Battery",
                     fontsize=13,fontweight="bold")
        solar_sweep_plot(sweep_df, sweep_bat_kwh, sweep_tariff, sweep_opt_kw, ax)
        plt.tight_layout(); save(fig,"08_solar_sweep.png")

    # 13 monthly savings — all quotes
    fig,ax = plt.subplots(figsize=(14,6))
    fig.suptitle("Expected Monthly Savings vs No-Solar Baseline",
                 fontsize=13,fontweight="bold")
    monthly_savings_plot(all_res, raw_df, ax)
    plt.tight_layout(); save(fig,"13_monthly_savings.png")

    # 09–12 representative day per season (energy flows + battery SoC)
    seasons = [
        ("09_summer_day.png", "Summer (Dec–Feb)", [12, 1, 2]),
        ("10_autumn_day.png", "Autumn (Mar–May)", [3, 4, 5]),
        ("11_winter_day.png", "Winter (Jun–Aug)", [6, 7, 8]),
        ("12_spring_day.png", "Spring (Sep–Nov)", [9, 10, 11]),
    ]
    for fname, label, months in seasons:
        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(13, 8),
                                        gridspec_kw={"height_ratios": [2, 1]})
        date = seasonal_day_plot(rms, months, ax1, ax2)
        if date:
            fig.suptitle(
                f"{label} — Representative Day ({date.strftime('%A %d %b %Y')})\n"
                f"{rms['solar_kw']}kW solar / {rms['bat_kwh']}kWh battery ({rms['tariff']})",
                fontsize=12, fontweight="bold")
        plt.tight_layout(); save(fig, fname)

    # 17 predicted monthly bill — all scenarios combined
    fig, ax = plt.subplots(figsize=(14, 6))
    fig.suptitle("Predicted Monthly Bill — All Quotes & Scenarios",
                 fontsize=13, fontweight="bold")
    monthly_bill_comparison_plot(all_res, arbit_res or [], raw_df, ax)
    plt.tight_layout(); save(fig, "17_monthly_bill_all_scenarios.png")

    # 15 battery grid arbitrage — average daily profile
    if arbit_res:
        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(13, 8),
                                        gridspec_kw={"height_ratios": [2, 1]})
        fig.suptitle(f"Battery Grid Arbitrage — {arbit_res[0]['bat_kwh']} kWh Battery, No Solar\n"
                     "Charges from cheap super off-peak grid power (9am–3pm), discharges at expensive peak (3–9pm)",
                     fontsize=12, fontweight="bold")
        arbitrage_day_plot(arbit_res[0], ax1, ax2)
        plt.tight_layout(); save(fig, "15_battery_arbitrage_profile.png")

    # 16 zero-net sizing
    if zero_net:
        fig, ax = plt.subplots(figsize=(10, 6))
        fig.suptitle("Zero-Net Sizing — Minimum Solar+Battery for $0 Annual Electricity Cost",
                     fontsize=13, fontweight="bold")
        zero_net_plot(zero_net, zero_net_tariff, ax)
        plt.tight_layout(); save(fig, "16_zero_net_sizing.png")

    print(f"\n  All charts → {out_dir.resolve()}")


# ─── BASELINE VERIFICATION ────────────────────────────────────────────────────
def print_baseline_check(raw_df, source_label="unknown"):
    """Print a step-by-step breakdown of the baseline bill calculation so the
    user can verify it against their actual Synergy bills."""
    n_rows  = len(raw_df)
    n_days  = int(raw_df["date"].nunique()) if "date" in raw_df.columns else n_rows // 48
    total_kwh = raw_df["consumption_kwh"].sum()
    avg_daily = total_kwh / n_days if n_days else 0

    # A1 Flat
    a1_import  = (raw_df["consumption_kwh"] * A1_RATE).sum()
    a1_supply  = A1_SUPPLY * n_days
    a1_export  = 0.0                          # no solar in base case
    a1_annual  = a1_import + a1_supply - a1_export

    # Midday Saver
    slots      = raw_df["slot"].values if "slot" in raw_df.columns else (
                    raw_df["datetime"].dt.hour * 2 + raw_df["datetime"].dt.minute // 30
                 ).values.astype(int)
    load       = raw_df["consumption_kwh"].values
    ms_import  = float(sum(load[i] * ms_rate(int(slots[i])) for i in range(len(raw_df))))
    ms_supply  = MS_SUPPLY * n_days
    ms_annual  = ms_import + ms_supply

    # Month-by-month breakdown (for cross-checking chart 14)
    rd = raw_df.copy()
    rd["month"] = rd["datetime"].dt.month
    rd["date"]  = rd["datetime"].dt.date
    rd["ic_a1"] = rd["consumption_kwh"] * A1_RATE
    rd["ic_ms"] = rd.apply(lambda r: r["consumption_kwh"] * ms_rate(int(r["slot"])), axis=1)
    mb = rd.groupby("month").agg(
        ic_a1=("ic_a1","sum"), ic_ms=("ic_ms","sum"),
        kwh=("consumption_kwh","sum"), days=("date","nunique"))
    mb["net_a1"] = mb["ic_a1"] + mb["days"] * A1_SUPPLY
    mb["net_ms"] = mb["ic_ms"] + mb["days"] * MS_SUPPLY
    month_names  = ["Jan","Feb","Mar","Apr","May","Jun",
                    "Jul","Aug","Sep","Oct","Nov","Dec"]

    B = "="*78; S = "-"*78
    print(f"\n{B}")
    print("  BASELINE BILL VERIFICATION (no solar / no battery)")
    print(B)
    print(f"  Data source  : {source_label}")
    print(f"  Data rows    : {n_rows:,}  ({n_rows/48:.1f} data-days equiv.)")
    print(f"  Unique dates : {n_days} days  "
          f"({'leap year — 366 days expected' if n_days==366 else '365 days expected' if n_days==365 else 'partial year'})")
    print(f"  Total load   : {total_kwh:,.1f} kWh  (avg {avg_daily:.2f} kWh/day)")
    print(S)
    print(f"  {'Item':<38} {'A1 Flat':>10}  {'Midday Saver':>12}")
    print(S)
    print(f"  {'Import cost (usage charges)':<38} ${a1_import:>9,.2f}  ${ms_import:>11,.2f}")
    print(f"  {'Supply charge':<38} ${a1_supply:>9,.2f}  ${ms_supply:>11,.2f}")
    print(f"    ({n_days} days × {A1_SUPPLY*100:.4f}c / {MS_SUPPLY*100:.4f}c per day)")
    print(f"  {'ANNUAL TOTAL':<38} ${a1_annual:>9,.2f}  ${ms_annual:>11,.2f}")
    print(f"  {'Average monthly total':<38} ${a1_annual/12:>9,.2f}  ${ms_annual/12:>11,.2f}")
    print(S)
    print(f"  MONTH-BY-MONTH (cross-check these against chart 14 base-case line):")
    print(f"  {'Month':<6} {'kWh':>7}  {'Days':>4}  {'A1 Net':>9}  {'MS Net':>9}")
    print(f"  {'-'*44}")
    for mo in range(1, 13):
        if mo in mb.index:
            row = mb.loc[mo]
            print(f"  {month_names[mo-1]:<6} {row['kwh']:>7.1f}  {int(row['days']):>4}  "
                  f"${row['net_a1']:>8,.2f}  ${row['net_ms']:>8,.2f}")
    print(S)
    print("  NOTE: These figures cover ONLY the main meter channel in the export")
    print("  file. Controlled-load tariffs (hot-water boost etc.) and any other")
    print("  separately metered circuits are NOT included. If your Synergy bill")
    print("  is higher, check whether you have a controlled-load tariff.")
    print(B)


# ─── SUMMARY PRINT ────────────────────────────────────────────────────────────
def print_summary(all_res, pbs, raw_df):
    B="="*78; S="-"*78
    base_a1 = baseline(raw_df,"A1 Flat"); base_ms = baseline(raw_df,"Midday Saver")
    total   = raw_df["consumption_kwh"].sum()
    avg_day = total / (len(raw_df)/48)
    eff_pct = (SHADE_SUMMER*3 + SHADE_AUTUMN*3 + SHADE_WINTER*3 + SHADE_SPRING*3) / 12 * 100

    print(f"\n{B}")
    print("  SOLAR & BATTERY ECONOMIC MODEL — Synergy WA | Rates: 1 Jul 2025")
    print(B)
    print(f"\n  ── Your Actual 2024 Consumption ──")
    print(f"  Annual total        : {total:>10,.1f} kWh")
    print(f"  Average daily       : {avg_day:>10.2f} kWh/day")
    print(f"  Baseline cost (A1)  : ${base_a1:>10,.2f}")
    print(f"  Baseline cost (MS)  : ${base_ms:>10,.2f}")
    print(f"\n  ── Solar Model Calibration ──")
    print(f"  Latitude            : {PERTH_LAT_DEG}° (Perth)  — sin(elevation) shape")
    print(f"  Unshaded yield      : {PERTH_SPECIFIC_YIELD:,} kWh/kWp/yr  (post SYS_EFF {SYS_EFF*100:.0f}%)")
    print(f"  Calibration factor  : {SOLAR_K:.4f}  (computed at startup)")
    print(f"\n  ── Shading Factors  (avg ~{eff_pct:.0f}% of unshaded) ──")
    print(f"  Summer (Dec–Feb): {SHADE_SUMMER*100:.0f}% | Autumn (Mar–May): {SHADE_AUTUMN*100:.0f}%")
    print(f"  Winter (Jun–Aug): {SHADE_WINTER*100:.0f}% | Spring (Sep–Nov): {SHADE_SPRING*100:.0f}%")
    print(f"  Effective site yield: ~{PERTH_SPECIFIC_YIELD*eff_pct/100:,.0f} kWh/kWp/yr with these shading factors")

    # Annual solar generation per year (degradation table) for unique solar sizes
    seen_kw = set()
    solar_quotes = [(r["label"], r["solar_kw"]) for r in all_res if r["solar_kw"] > 0 and r["solar_kw"] not in seen_kw and not seen_kw.add(r["solar_kw"])]
    if solar_quotes:
        print(f"\n  ── Annual Solar Generation by Year (kWh) ──")
        yrs = [1, 5, 10, 15, 20, 25]
        hdr = f"  {'System':<28} {'kW':>5} {'kWh/kWp':>8}  " + "  ".join(f"Yr{y:>2}" for y in yrs)
        print(hdr); print(f"  {'-'*len(hdr.expandtabs())}")
        for lbl, kw in solar_quotes:
            yr1 = next(r["annual_solar_kwh"] for r in all_res if r["solar_kw"]==kw)
            spec = yr1 / kw if kw else 0
            gen_by_yr = [f"{yr1*(1-SOL_DEG)**(y-1):>6,.0f}" for y in yrs]
            print(f"  {lbl[:27]:<28} {kw:>5.1f} {spec:>8,.0f}  {'  '.join(gen_by_yr)}")
        print(f"  (Perth PVWatts reference: ~1,550–1,750 kWh/kWp/yr unshaded; typical shaded site ~1,100–1,400)")

    print(f"\n  ── Year-1 Results ──")
    print(f"  {'System':<36} {'Tariff':<15} {'Solar':>8} {'Import':>8} {'Export':>8} {'Net Cost':>10} {'Saving':>9} {'Self-suf':>9}")
    print(f"  {S}")
    for r in all_res:
        base = base_a1 if r["tariff"]=="A1 Flat" else base_ms
        sav  = base - r["net_cost"]
        print(f"  {r['label'][:35]:<36} {r['tariff']:<15}"
              f"  {r['annual_solar_kwh']:>6,.0f}kWh"
              f"  {r['annual_import_kwh']:>6,.0f}kWh"
              f"  {r['annual_export_kwh']:>6,.0f}kWh"
              f"  ${r['net_cost']:>8,.2f}"
              f"  ${sav:>7,.2f}"
              f"  {r['self_suf_pct']:>7.1f}%")
    print(f"\n{B}")
    print("  PAYBACK & FINANCIAL SUMMARY (25-Year Horizon)")
    print(B)
    print(f"  {'Label':<36} {'Tariff':<15} {'Gross':>9} {'Net':>9} {'Payback':>8} {'Disc Pb':>8} {'NPV@8%':>10} {'25yr Save':>11} {'ROI':>8}")
    print(f"  {S}")
    for pb in pbs:
        pbl  = f"{pb['pb_yr']}y"       if pb["pb_yr"]       else ">25y"
        pbld = f"{pb['pb_yr_disc']}y"  if pb.get("pb_yr_disc") else ">25y"
        npv_s = f"${pb['npv']:>8,.0f}" if pb.get("npv") is not None else "       n/a"
        print(f"  {pb['label'][:35]:<36} {pb['tariff']:<15}"
              f"  ${pb['cost']:>7,.0f}"
              f"  ${pb['net']:>7,.0f}"
              f"  {pbl:>7}"
              f"  {pbld:>7}"
              f"  {npv_s:>10}"
              f"  ${pb['total_save']:>9,.0f}"
              f"  {pb['roi']:>7.1f}%")
    print(f"\n  ── Rebates ──")
    seen=set()
    for pb in pbs:
        if pb["label"] not in seen:
            seen.add(pb["label"])
            print(f"  {pb['label']}")
            print(f"    Federal : ${pb['fed']:>8,.0f}  (tiered: $252/kWh 0–14kWh, $151/kWh 14–28kWh — from 1 May 2026)")
            print(f"    State   : ${pb['state']:>8,.0f}  (Synergy WA Residential Battery Scheme, cap $1,300)")
            print(f"    Net out : ${pb['net']:>8,.0f}")
    print(f"\n{B}\n")


# ─── DATA LOADING ─────────────────────────────────────────────────────────────
def load_data(path):
    # Validate file path before attempting to read
    p = validate_file_path(path)
    suffix = p.suffix.lower()

    try:
        if suffix in (".xlsx", ".xls"):
            raw = pd.read_excel(p)
        elif suffix == ".tsv":
            raw = pd.read_csv(p, sep="\t")
        else:
            raw = pd.read_csv(p)
    except Exception as exc:
        raise ValidationError(f"Failed to read '{p}': {exc}")

    if raw.empty:
        raise ValidationError(f"File loaded but contains no data rows: {p}")

    raw.columns = raw.columns.str.strip()

    # Detect column structure
    try:
        if "Date" in raw.columns and "Time" in raw.columns:
            raw["datetime"] = pd.to_datetime(
                raw["Date"].astype(str)+" "+raw["Time"].astype(str), dayfirst=True,
                errors="coerce"
            )
            kwh_col = next((c for c in raw.columns
                            if c.lower() in ("kwh","consumption","usage","energy")), None)
            if not kwh_col:
                raise ValidationError(
                    f"Found Date+Time columns but no kWh/consumption column. "
                    f"Available columns: {list(raw.columns)}"
                )
            raw["consumption_kwh"] = pd.to_numeric(raw[kwh_col], errors="coerce")
        else:
            raw.columns = raw.columns.str.lower().str.replace(" ","_")
            dc = next((c for c in raw.columns if "date" in c or "time" in c), None)
            kc = next((c for c in raw.columns if any(k in c for k in
                       ["kwh","consumption","usage","energy"])), None)
            if not dc:
                raise ValidationError(
                    f"Could not detect a datetime column. "
                    f"Expected one of: datetime, timestamp, date_time, or Date+Time. "
                    f"Found: {list(raw.columns)}"
                )
            if not kc:
                raise ValidationError(
                    f"Could not detect a consumption column. "
                    f"Expected one of: kwh, consumption, usage, energy. "
                    f"Found: {list(raw.columns)}"
                )
            raw["datetime"] = pd.to_datetime(raw[dc], dayfirst=True, errors="coerce")
            raw["consumption_kwh"] = pd.to_numeric(raw[kc], errors="coerce")
    except ValidationError:
        raise
    except Exception as exc:
        raise ValidationError(f"Error parsing data columns: {exc}")

    # Check parse failures
    n_bad_dt = raw["datetime"].isna().sum()
    if n_bad_dt > 0:
        _warn(f"{n_bad_dt} rows had unparseable datetime values — dropping them.")
        raw = raw[raw["datetime"].notna()].reset_index(drop=True)

    if len(raw) == 0:
        raise ValidationError("All rows had unparseable datetime values.")

    df = raw[["datetime","consumption_kwh"]].sort_values("datetime").reset_index(drop=True)

    # Run the comprehensive validator
    df = validate_consumption_data(df)

    print(f"  Loaded {len(df):,} rows | {df['consumption_kwh'].sum():.1f} kWh | "
          f"{df['datetime'].iloc[0].date()} → {df['datetime'].iloc[-1].date()}")
    return df


def demo_data():
    print("  Generating synthetic data (~24 kWh/day Perth household).")
    rng = np.random.default_rng(42)
    idx = pd.date_range("2024-01-01", periods=48*366, freq="30min")
    h   = idx.hour + idx.minute/60
    doy = idx.dayofyear
    sa  = 2*np.pi*(doy-355)/365   # Southern Hemisphere: summer peaks Dec/Jan
    profile = 0.20 + 0.30*np.exp(-((h-7.5)**2)/2) + \
              0.55*np.exp(-((h-18.5)**2)/4) + \
              0.5*np.clip(np.cos(sa),0,1)*0.35*np.exp(-((h-14)**2)/8)
    kwh = np.clip(profile*rng.lognormal(0,0.2,len(idx))*0.5, 0.02, 5.0)
    return pd.DataFrame({"datetime":idx,"consumption_kwh":kwh})


# ─── EXCEL EXPORT ────────────────────────────────────────────────────────────
def write_excel(all_res, pbs, raw_df, sweep_df, opt_kw, out_dir):
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter

    path     = out_dir / "solar_analysis.xlsx"
    base_a1  = baseline(raw_df, "A1 Flat")
    base_ms  = baseline(raw_df, "Midday Saver")
    total    = raw_df["consumption_kwh"].sum()
    span     = (raw_df["datetime"].iloc[-1] - raw_df["datetime"].iloc[0]).days + 1

    def style_header(ws, row=1, fill="1a73e8"):
        fill_obj = PatternFill("solid", fgColor=fill)
        for cell in ws[row]:
            if cell.value is not None:
                cell.font      = Font(bold=True, color="FFFFFF")
                cell.fill      = fill_obj
                cell.alignment = Alignment(horizontal="center", wrap_text=True)

    def autofit(ws, min_w=10, max_w=40):
        for col in ws.columns:
            width = max(len(str(c.value or "")) for c in col)
            ws.column_dimensions[get_column_letter(col[0].column)].width = \
                min(max(width + 2, min_w), max_w)

    with pd.ExcelWriter(path, engine="openpyxl") as writer:

        # ── Sheet 1: Summary
        rows = [
            ["CONSUMPTION"],
            ["Annual total (kWh)",              f"{total:,.1f}"],
            ["Average daily (kWh/day)",          f"{total/span:.2f}"],
            ["Data span (days)",                 span],
            [],
            ["BASELINE COSTS (no solar/battery)"],
            ["A1 Flat annual cost ($)",          f"${base_a1:,.2f}"],
            ["Midday Saver annual cost ($)",     f"${base_ms:,.2f}"],
            [],
            ["SHADING MODEL"],
            ["Annual effective output",          "~40% of unshaded"],
            ["Summer (Dec–Feb)",                 "75%"],
            ["Autumn (Mar–May)",                 "55%"],
            ["Winter (Jun–Aug)",                 "30%"],
            ["Spring (Sep–Nov)",                 "50%"],
            [],
            ["REBATES (1 May 2026)"],
            ["Federal tier 1 (0–14 kWh)",        "$252/kWh"],
            ["Federal tier 2 (14–28 kWh)",       "$151/kWh"],
            ["Federal tier 3 (28–50 kWh)",       "$38/kWh"],
            ["State (Synergy WA, VPP required)", f"${STATE_REBATE_FLAT:,.0f} flat"],
            [],
            ["ANALYSIS PARAMETERS"],
            ["Analysis horizon (years)",         ANALYSIS_YEARS],
            ["Tariff escalation p.a.",           f"{TARIFF_ESC*100:.1f}%"],
            ["DEBS decline p.a.",                f"{DEBS_DECL*100:.1f}%"],
            ["Battery round-trip efficiency",    f"{BAT_RTE*100:.0f}%"],
            ["Battery depth of discharge",       f"{BAT_DOD*100:.0f}%"],
            ["Battery degradation p.a.",         f"{BAT_DEG*100:.1f}%"],
            ["Panel degradation p.a.",           f"{SOL_DEG*100:.2f}%"],
            ["System efficiency (inv+wiring)",   f"{SYS_EFF*100:.0f}%"],
        ]
        pd.DataFrame(rows).to_excel(writer, sheet_name="Summary", index=False, header=False)
        ws = writer.sheets["Summary"]
        for row in ws.iter_rows():
            if row[0].value and str(row[0].value).isupper():
                row[0].font = Font(bold=True)
        autofit(ws)

        # ── Sheet 2: Year-1 Results
        yr1_rows = []
        for r in all_res:
            base = base_a1 if r["tariff"] == "A1 Flat" else base_ms
            yr1_rows.append({
                "Label":               r["label"],
                "Tariff":              r["tariff"],
                "Solar (kW)":          r["solar_kw"],
                "Battery (kWh)":       r["bat_kwh"],
                "Solar Gen (kWh)":     round(r["annual_solar_kwh"], 0),
                "Grid Import (kWh)":   round(r["annual_import_kwh"], 0),
                "Grid Export (kWh)":   round(r["annual_export_kwh"], 0),
                "Import Cost ($)":     round(r["import_cost"], 2),
                "Supply Cost ($)":     round(r["supply_cost"], 2),
                "Export Credit ($)":   round(r["export_credit"], 2),
                "Net Annual Cost ($)": round(r["net_cost"], 2),
                "Annual Saving ($)":   round(base - r["net_cost"], 2),
                "Self-sufficiency (%)":round(r["self_suf_pct"], 1),
            })
        pd.DataFrame(yr1_rows).to_excel(writer, sheet_name="Year-1 Results", index=False)
        ws = writer.sheets["Year-1 Results"]
        style_header(ws); autofit(ws)
        ws.freeze_panes = "A2"

        # ── Sheet 3: Payback — cumulative cashflow table
        pb_cf = {"Year": list(range(ANALYSIS_YEARS + 1))}
        for pb in pbs:
            pb_cf[f"{pb['label'][:28]} | {pb['tariff']}"] = pb["cum"]
        pd.DataFrame(pb_cf).to_excel(writer, sheet_name="25-Year Cashflow", index=False)
        ws = writer.sheets["25-Year Cashflow"]
        style_header(ws); autofit(ws)
        ws.freeze_panes = "B2"

        # ── Sheet 4: Payback Summary
        pb_rows = []
        for pb in pbs:
            pb_rows.append({
                "Label":               pb["label"],
                "Tariff":              pb["tariff"],
                "Solar (kW)":          pb["solar_kw"],
                "Battery (kWh)":       pb["bat_kwh"],
                "Gross Cost ($)":      pb["cost"],
                "Federal Rebate ($)":  round(pb["fed"], 0),
                "State Rebate ($)":    pb["state"],
                "Net Out-of-Pocket ($)":round(pb["net"], 0),
                "Payback (years)":          pb["pb_yr"] if pb["pb_yr"] else ">25",
                "Opp-Cost Payback (years)": pb.get("pb_yr_disc") or ">25",
                "NPV @ 8% ($)":            round(pb.get("npv", 0), 0),
                "25-yr Total Saving ($)":  round(pb["total_save"], 0),
                "25-yr ROI (%)":           round(pb["roi"], 1),
                "Year-1 Saving ($)":       round(pb["yr1_save"], 0),
            })
        pd.DataFrame(pb_rows).to_excel(writer, sheet_name="Payback Summary", index=False)
        ws = writer.sheets["Payback Summary"]
        style_header(ws); autofit(ws)
        ws.freeze_panes = "A2"

        # ── Sheet 5: Monthly Savings
        month_lbls = ["Jan","Feb","Mar","Apr","May","Jun",
                      "Jul","Aug","Sep","Oct","Nov","Dec"]
        months = list(range(1, 13))
        rd = raw_df.copy()
        rd["month"] = rd["datetime"].dt.month
        rd["date"]  = rd["datetime"].dt.date
        rd["ic_a1"] = rd["consumption_kwh"] * A1_RATE
        rd["ic_ms"] = rd.apply(lambda r: r["consumption_kwh"] * ms_rate(int(r["slot"])), axis=1)
        mb = rd.groupby("month").agg(
            ic_a1=("ic_a1","sum"), ic_ms=("ic_ms","sum"), days=("date","nunique"))
        mb = mb.reindex(months, fill_value=0)
        mb["base_a1"] = mb["ic_a1"] + mb["days"] * A1_SUPPLY
        mb["base_ms"] = mb["ic_ms"] + mb["days"] * MS_SUPPLY
        ms_rows = {"Month": month_lbls,
                   "Baseline A1 ($)":     mb["base_a1"].round(2).tolist(),
                   "Baseline Midday ($)": mb["base_ms"].round(2).tolist()}
        for r in all_res:
            df = r["df"].copy()
            df["month"] = df["datetime"].dt.month
            df["date"]  = df["datetime"].dt.date
            if r["tariff"] == "A1 Flat":
                df["ic"] = df["g_imp"] * A1_RATE
                sc, base_m = A1_SUPPLY, mb["base_a1"]
            else:
                df["ic"] = df.apply(
                    lambda row: row["g_imp"] * ms_rate(int(row["slot"])), axis=1)
                sc, base_m = MS_SUPPLY, mb["base_ms"]
            df["ec"] = df.apply(
                lambda row: row["g_exp"] * debs_rate(int(row["slot"])), axis=1)
            m = df.groupby("month").agg(
                ic=("ic","sum"), ec=("ec","sum"), days=("date","nunique"))
            m = m.reindex(months, fill_value=0)
            m["net"]    = m["ic"] + m["days"] * sc - m["ec"]
            col_name    = f"{r['label'][:28]} | {r['tariff']}"
            ms_rows[col_name] = (base_m.values - m["net"].values).round(2).tolist()
        pd.DataFrame(ms_rows).to_excel(
            writer, sheet_name="Monthly Savings", index=False)
        ws = writer.sheets["Monthly Savings"]
        style_header(ws); autofit(ws)
        ws.freeze_panes = "B2"

        # ── Sheet 5b: Monthly Costs (all quotes — for bill verification)
        mc_rows = {"Month": month_lbls}
        for r in all_res:
            df = r["df"].copy()
            df["month"] = r["df"]["datetime"].dt.month
            df["date"]  = r["df"]["datetime"].dt.date
            if r["tariff"] == "A1 Flat":
                df["ic"] = df["g_imp"] * A1_RATE
                sc2 = A1_SUPPLY
            else:
                df["ic"] = df.apply(
                    lambda row: row["g_imp"] * ms_rate(int(row["slot"])), axis=1)
                sc2 = MS_SUPPLY
            df["ec"] = df.apply(
                lambda row: row["g_exp"] * debs_rate(int(row["slot"])), axis=1)
            m2 = df.groupby("month").agg(
                ic=("ic","sum"), ec=("ec","sum"), days=("date","nunique"))
            m2 = m2.reindex(months, fill_value=0)
            m2["net"] = m2["ic"] + m2["days"] * sc2 - m2["ec"]
            col_name = f"{r['label'][:28]} | {r['tariff']}"
            mc_rows[col_name] = m2["net"].round(2).tolist()
        pd.DataFrame(mc_rows).to_excel(
            writer, sheet_name="Monthly Costs (All)", index=False)
        ws = writer.sheets["Monthly Costs (All)"]
        style_header(ws); autofit(ws)
        ws.freeze_panes = "B2"

        # ── Sheet 5: Solar Sweep
        if sweep_df is not None:
            sweep_out = sweep_df.rename(columns={
                "solar_kw":      "Solar (kW)",
                "inv_kw":        "Inverter (kW)",
                "self_suf_pct":  "Self-sufficiency (%)",
                "bat_util_pct":  "Battery utilisation (%)",
                "days_full_pct": "Days battery full (%)",
                "export_kwh":    "Annual export (kWh)",
                "import_kwh":    "Annual import (kWh)",
                "net_cost":      "Net annual cost ($)",
                "annual_savings":"Annual saving ($)",
            }).copy()
            sweep_out["Recommended"] = sweep_out["Solar (kW)"].apply(
                lambda x: "YES" if x == opt_kw else "")
            sweep_out.to_excel(writer, sheet_name="Solar Sweep", index=False)
            ws = writer.sheets["Solar Sweep"]
            style_header(ws); autofit(ws)
            ws.freeze_panes = "A2"
            # Highlight recommended row
            highlight = PatternFill("solid", fgColor="FFF3CD")
            for row in ws.iter_rows(min_row=2):
                if row[-1].value == "YES":
                    for cell in row:
                        cell.fill = highlight

    print(f"  Excel        → {path.resolve()}")
    return path


# ─── LARGEST SOLAR QUOTE HALF-HOURLY TIME-SERIES EXPORT ──────────────────────
def write_timeseries_excel(raw_df, out_dir):
    """Export a full half-hourly time series for the largest solar+battery quote.

    Sheets:
      'Midday Saver'  — half-hourly energy flows + per-slot costs on MS tariff
      'A1 Flat'       — same on A1 Flat tariff
      'Daily Summary' — day-level aggregation of both tariffs
    """
    try:
        from openpyxl.styles import Font, PatternFill, Alignment
        from openpyxl.utils import get_column_letter
    except ImportError:
        _warn("openpyxl not installed — skipping time-series Excel export.")
        return

    # Pick the quote with the largest solar array (skip base case which has 0 kW)
    solar_quotes = [(q, q[1]) for q in QUOTES if q[1] > 0]
    if not solar_quotes:
        _warn("No solar quotes found — skipping time-series export.")
        return
    selected_quote = max(solar_quotes, key=lambda x: x[1])[0]
    ts_label = selected_quote[0]
    _, solar_kw, bat_kwh, inv_kw, _ = selected_quote
    df_solar = add_solar(raw_df, solar_kw)

    def style_header(ws, fill_hex="1a73e8"):
        fill = PatternFill("solid", fgColor=fill_hex)
        for cell in ws[1]:
            if cell.value is not None:
                cell.font      = Font(bold=True, color="FFFFFF")
                cell.fill      = fill
                cell.alignment = Alignment(horizontal="center", wrap_text=True)

    def autofit(ws, min_w=10, max_w=30):
        for col in ws.columns:
            width = max(len(str(c.value or "")) for c in col)
            ws.column_dimensions[
                get_column_letter(col[0].column)
            ].width = min(max(width + 2, min_w), max_w)

    path = out_dir / "timeseries_largest_solar.xlsx"
    with pd.ExcelWriter(path, engine="openpyxl") as writer:

        for tariff in ["Midday Saver", "A1 Flat"]:
            res     = simulate(df_solar, solar_kw, bat_kwh, inv_kw, tariff, yr_offset=0)
            sim_df  = res["df"].copy()
            slots   = sim_df["slot"].values.astype(int)

            if tariff == "Midday Saver":
                rates = np.vectorize(ms_rate)(slots)
            else:
                rates = np.full(len(sim_df), A1_RATE)

            debs = np.vectorize(debs_rate)(slots)

            out = pd.DataFrame({
                "Datetime":              sim_df["datetime"],
                "Date":                  sim_df["datetime"].dt.date,
                "Time":                  sim_df["datetime"].dt.strftime("%H:%M"),
                "Slot":                  slots,
                "Load (kWh)":            sim_df["consumption_kwh"].round(4),
                "Solar Gen (kWh)":       sim_df["solar_kwh"].round(4),
                "Solar Self-Use (kWh)":  sim_df["s_slf"].round(4),
                "Bat Charge (kWh)":      sim_df["b_chg"].round(4),
                "Bat Discharge (kWh)":   (sim_df["b_dis"] * BAT_RTE**0.5).round(4),
                "Grid Import (kWh)":     sim_df["g_imp"].round(4),
                "Grid Export (kWh)":     sim_df["g_exp"].round(4),
                "Battery SoC (kWh)":     sim_df["soc"].round(4),
                "Tariff Rate ($/kWh)":   np.round(rates, 6),
                "Import Cost ($)":       np.round(sim_df["g_imp"].values * rates, 6),
                "Export Credit ($)":     np.round(sim_df["g_exp"].values * debs, 6),
                "Net Slot Cost ($)":     np.round(
                    sim_df["g_imp"].values * rates - sim_df["g_exp"].values * debs, 6),
            })

            sheet_name = tariff
            out.to_excel(writer, sheet_name=sheet_name, index=False)
            ws = writer.sheets[sheet_name]
            style_header(ws, "2c5f8a" if tariff == "Midday Saver" else "1a73e8")
            autofit(ws)
            ws.freeze_panes = "A2"

        # ── Daily Summary sheet (both tariffs side by side) ───────────────────
        daily_rows = []
        for tariff in ["Midday Saver", "A1 Flat"]:
            res    = simulate(df_solar, solar_kw, bat_kwh, inv_kw, tariff, yr_offset=0)
            sim_df = res["df"].copy()
            sim_df["_date"] = sim_df["datetime"].dt.date
            grp = sim_df.groupby("_date")
            for date, g in grp:
                if tariff == "Midday Saver":
                    day_rates = np.vectorize(ms_rate)(g["slot"].values.astype(int))
                else:
                    day_rates = np.full(len(g), A1_RATE)
                daily_rows.append({
                    "Date":                  date,
                    "Tariff":                tariff,
                    "Load (kWh)":            round(g["consumption_kwh"].sum(), 3),
                    "Solar Gen (kWh)":       round(g["solar_kwh"].sum(), 3),
                    "Solar Self-Use (kWh)":  round(g["s_slf"].sum(), 3),
                    "Bat Discharge (kWh)":   round((g["b_dis"] * BAT_RTE**0.5).sum(), 3),
                    "Grid Import (kWh)":     round(g["g_imp"].sum(), 3),
                    "Grid Export (kWh)":     round(g["g_exp"].sum(), 3),
                    "Import Cost ($)":       round((g["g_imp"].values * day_rates).sum(), 4),
                    "Export Credit ($)":     round(
                        (g["g_exp"].values * np.vectorize(debs_rate)(
                            g["slot"].values.astype(int))).sum(), 4),
                    "Self-Suf (%)":          round(
                        (g["s_slf"].sum() + (g["b_dis"] * BAT_RTE**0.5).sum())
                        / max(g["consumption_kwh"].sum(), 1e-6) * 100, 1),
                })

        daily_df = pd.DataFrame(daily_rows).sort_values(["Date", "Tariff"])
        daily_df.to_excel(writer, sheet_name="Daily Summary", index=False)
        ws = writer.sheets["Daily Summary"]
        style_header(ws, "2d6a4f")
        autofit(ws)
        ws.freeze_panes = "A2"

    print(f"  Timeseries ({ts_label[:30]}) → {path.resolve()}")
    return path


# ─── WORD REPORT ──────────────────────────────────────────────────────────────
def write_report(all_res, pbs, raw_df, sweep_df, opt_kw, out_dir,
                 arbit_res=None, arbit_pbs=None, zero_net=None):
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        _warn("python-docx not installed — skipping Word report.\n"
              "    Install with: pip install python-docx")
        return

    import datetime as dt

    base_a1 = baseline(raw_df, "A1 Flat")
    base_ms = baseline(raw_df, "Midday Saver")
    total   = raw_df["consumption_kwh"].sum()
    span    = (raw_df["datetime"].iloc[-1] - raw_df["datetime"].iloc[0]).days + 1

    doc = Document()

    def bold_run(para, text):
        run = para.add_run(text)
        run.bold = True
        return run

    def add_table(doc, df):
        tbl = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
        tbl.style = "Table Grid"
        for j, col in enumerate(df.columns):
            cell = tbl.rows[0].cells[j]
            cell.text = col
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, (_, row) in enumerate(df.iterrows()):
            for j, val in enumerate(row):
                cell = tbl.rows[i + 1].cells[j]
                cell.text = str(val)
                cell.paragraphs[0].runs[0].font.size = Pt(8)
        return tbl

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_heading("Solar & Battery Economic Analysis", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        f"Synergy WA (Perth / SWIS)  |  Generated {dt.date.today().strftime('%d %B %Y')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── How to read this report ───────────────────────────────────────────────
    doc.add_heading("How to Read This Report", 1)
    doc.add_paragraph(
        "This report models the financial impact of solar and battery storage at your property, "
        "using your actual Synergy half-hourly smart meter data. Sections 1–5 explain the "
        "modelling methodology; Sections 6–7 give year-1 and 25-year financial results per quote; "
        "Sections 8–9 cover zero-net sizing and solar optimisation; "
        "Section 10 lists model caveats; Section 11 contains supporting charts. "
        "Check Fig 5 first — if the modelled baseline deviates significantly from your actual "
        "Synergy bills, the savings projections will be proportionally off."
    )

    # ── Executive Summary ─────────────────────────────────────────────────────
    doc.add_heading("Executive Summary", 1)
    best_pb = min((pb for pb in pbs if pb["pb_yr"]), key=lambda x: x["pb_yr"], default=None)
    doc.add_paragraph(
        f"This report evaluates the financial case for installing rooftop solar and/or a battery "
        f"storage system at your property, based on your actual 2024 Synergy smart meter data. "
        f"Your household consumed {total:,.0f} kWh over {span} days "
        f"({total/span:.1f} kWh/day average). On today's standard A1 Flat tariff, that consumption "
        f"costs approximately ${base_a1:,.0f} per year (${base_a1/12:,.0f}/month). On the Midday "
        f"Saver time-of-use tariff — which has much cheaper rates during the middle of the day and "
        f"much more expensive rates in the evening — your current usage pattern costs "
        f"approximately ${base_ms:,.0f} per year before any solar or battery is installed. "
        f"Whether Midday Saver is cheaper than A1 Flat for you without solar depends heavily on "
        f"when you use electricity; adding solar almost always makes Midday Saver the better choice."
    )
    if best_pb:
        doc.add_paragraph(
            f"Across all quotes modelled, the fastest payback is '{best_pb['label']}' on the "
            f"{best_pb['tariff']} tariff, with an estimated payback period of {best_pb['pb_yr']} "
            f"years and a cumulative saving of ${best_pb['total_save']:,.0f} over 25 years after "
            f"deducting rebates. This is a nominal figure — it does not discount future cash flows "
            f"to present value. The payback period is the most useful single metric for comparing "
            f"options: shorter payback means less financial risk and faster recovery of your "
            f"investment."
        )
    doc.add_paragraph(
        "All figures are model estimates — actual performance will depend on weather, shading, "
        "installation quality, and future tariff changes. Use this report as a decision-support "
        "tool alongside installer quotes, not as a guarantee of financial outcome."
    )

    # ── 1. Consumption ────────────────────────────────────────────────────────
    doc.add_heading("1. Your Electricity Consumption", 1)
    doc.add_paragraph(
        f"The model uses all {span*48:,} half-hour consumption slots in your {span}-day dataset "
        "without averaging. The modelled baseline bill should closely match your actual Synergy "
        "statements (verify using Fig 5). Solar generation is highest in summer, offsetting "
        "your highest cooling bills; winter heating load falls mostly in the evening when solar "
        "is zero — which is why a battery is especially valuable in Perth."
    )
    p = doc.add_paragraph()
    bold_run(p, "Annual total: ");           p.add_run(f"{total:,.0f} kWh\n")
    bold_run(p, "Average daily: ");          p.add_run(f"{total/span:.1f} kWh/day\n")
    bold_run(p, "Data period: ");            p.add_run(f"{span} days\n")
    bold_run(p, "Baseline — A1 Flat: ");     p.add_run(f"${base_a1:,.2f}/year  (${base_a1/12:,.2f}/month)\n")
    bold_run(p, "Baseline — Midday Saver: ");p.add_run(f"${base_ms:,.2f}/year  (${base_ms/12:,.2f}/month)")
    doc.add_paragraph(
        "Note: controlled-load circuits (hot-water, pool pump) are sometimes metered separately "
        "and may not appear in the export. If your actual Synergy bills are higher than the "
        "modelled baseline, the solar savings will be proportionally understated."
    )

    # ── 2. How Solar Generation is Modelled ──────────────────────────────────
    doc.add_heading("2. How Solar Generation is Modelled", 1)
    doc.add_paragraph(
        "Because you do not yet have solar panels, the model constructs a synthetic solar output "
        f"profile for each day of the year. It uses the physical sin(elevation angle) shape for "
        f"Perth's latitude (−31.95°), calibrated to {PERTH_SPECIFIC_YIELD:,} kWh/kWp/year "
        f"(unshaded, post {SYS_EFF*100:.0f}% system efficiency). Each half-hour slot's output is "
        "reduced by the seasonal shading factors (Section 3) and a panel degradation rate of "
        f"{SOL_DEG*100:.2f}%/year. Surplus solar charges the battery first; any remaining export "
        "earns DEBS credits. When consumption exceeds solar the battery discharges before "
        "grid power is drawn."
    )

    # ── 3. Site Shading Model ─────────────────────────────────────────────────
    doc.add_heading("3. Site Shading Model", 1)
    doc.add_paragraph(
        "Shading is the largest source of uncertainty in any solar model. The sin(elevation) "
        "profile used here naturally goes to zero at sunrise/sunset; the seasonal factors below "
        "then scale the whole day's output to account for trees, neighbouring structures, and "
        "roof obstructions. An unshaded north-facing roof would use factors of 80–90% (summer) "
        "and 60–70% (winter). Partial shading on a string-inverter system is especially "
        "costly — consider optimisers or microinverters if your roof has shadows."
    )
    doc.add_paragraph(
        "The seasonal factors currently set in this model are:"
    )
    shading_df = pd.DataFrame({
        "Season":           ["Summer (Dec–Feb)", "Autumn (Mar–May)", "Winter (Jun–Aug)", "Spring (Sep–Nov)"],
        "Shading factor":   [f"{SHADE_SUMMER*100:.0f}%", f"{SHADE_AUTUMN*100:.0f}%",
                             f"{SHADE_WINTER*100:.0f}%", f"{SHADE_SPRING*100:.0f}%"],
        "Interpretation":   [
            "Sun is high; most shading at day edges only",
            "Transitional; moderate shading across the day",
            "Sun is low; broader shadows throughout the day",
            "Recovering toward summer levels",
        ],
    })
    add_table(doc, shading_df)
    doc.add_paragraph(
        "A ±10 percentage-point change in shading factors shifts annual solar yield by roughly "
        "±5–8%, with a similar effect on year-1 savings. If you obtain a professional shading "
        "assessment, update the SHADE_* constants at the top of the script and re-run."
    )

    # ── 4. How the Battery is Modelled ───────────────────────────────────────
    doc.add_heading("4. How the Battery is Modelled", 1)
    doc.add_paragraph(
        "Each half-hour: solar surplus charges the battery first; any remaining surplus exports "
        "to the grid. If consumption exceeds solar, the battery discharges before grid power is "
        "drawn. The battery never grid-charges in solar scenarios (only in the arbitrage scenario, "
        "Section 8). Key parameters:"
    )
    for bullet in [
        f"Round-trip efficiency: {BAT_RTE*100:.0f}% (√ applied symmetrically to charge and "
        f"discharge). Storing 10 kWh and retrieving it yields {BAT_RTE*10:.1f} kWh.",
        f"Depth of discharge: {BAT_DOD*100:.0f}% of nameplate — e.g. a 13.5 kWh battery "
        f"has {13.5*BAT_DOD:.1f} kWh usable capacity.",
        f"Annual capacity degradation: {BAT_DEG*100:.1f}%/year.",
        "Inverter power cap: charge/discharge rate is limited to the quoted inverter size (kW).",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    # ── 5. Rebates ────────────────────────────────────────────────────────────
    doc.add_heading("5. Applicable Rebates and Incentives", 1)
    doc.add_paragraph(
        "Two rebates apply to battery installations in WA as of May 2026, both reflected in "
        "the payback calculations:"
    )
    p = doc.add_paragraph()
    bold_run(p, "Federal — Cheaper Home Batteries Program (from 1 May 2026): ")
    p.add_run(
        f"tiered at $252/kWh (first 14 kWh), $151/kWh (14–28 kWh), $38/kWh (28–50 kWh), "
        "capped at 40% of installed system cost. Applied by the installer as a point-of-sale "
        "discount on the battery only.\n"
    )
    bold_run(p, f"State — Synergy WA Residential Battery Scheme: ")
    p.add_run(
        f"${STATE_REBATE_FLAT:,.0f} flat rebate, requires VPP enrolment (Synergy can dispatch "
        "small amounts of stored energy during grid stress events — infrequent in practice). "
        "All battery quotes assumed VPP-eligible.\n"
    )
    bold_run(p, "Solar STCs: ")
    p.add_run(
        "assumed already reflected in the installer's quoted price. Not modelled separately."
    )

    # ── 6. Year-1 Quote Comparison ────────────────────────────────────────────
    doc.add_heading("6. Year-1 Results by Quote", 1)
    doc.add_paragraph(
        "Year-1 modelled performance per quote — before tariff escalation, degradation, or DEBS "
        "rate changes. Saving = baseline − net cost. Self-suf = % of load met by solar+battery "
        "(not grid). Negative net cost = Synergy owes you a credit."
    )
    yr1_rows = []
    for r in all_res:
        base = base_a1 if r["tariff"] == "A1 Flat" else base_ms
        yr1_rows.append({
            "Quote":      r["label"][:30],
            "Tariff":     r["tariff"],
            "Solar kW":   r["solar_kw"],
            "Solar kWh":  f"{r['annual_solar_kwh']:,.0f}",
            "Import kWh": f"{r['annual_import_kwh']:,.0f}",
            "Export kWh": f"{r['annual_export_kwh']:,.0f}",
            "Net Cost":   f"${r['net_cost']:,.0f}",
            "Saving":     f"${base - r['net_cost']:,.0f}",
            "Self-suf":   f"{r['self_suf_pct']:.0f}%",
        })
    add_table(doc, pd.DataFrame(yr1_rows))
    doc.add_paragraph(
        "Note: A1 Flat and Midday Saver baselines differ (the same consumption costs a different "
        "amount on each tariff), so savings are always shown relative to the baseline on the "
        "same tariff. To compare a Midday Saver result directly to your current A1 Flat bill, "
        f"note that switching to Midday Saver without any solar costs approximately "
        f"${base_ms:,.0f}/year vs your current ${base_a1:,.0f}/year on A1 Flat — a difference "
        f"of ${abs(base_ms-base_a1):,.0f}/year based on your 2024 usage pattern."
    )

    # ── 7. Financial Summary (25 years) ──────────────────────────────────────
    doc.add_heading(f"7. Financial Summary — {ANALYSIS_YEARS}-Year Horizon", 1)
    doc.add_paragraph(
        f"Annual adjustments over the {ANALYSIS_YEARS}-year horizon: tariffs escalate "
        f"{TARIFF_ESC*100:.1f}%/yr; DEBS rates decline {DEBS_DECL*100:.0f}%/yr; panel output "
        f"degrades {SOL_DEG*100:.2f}%/yr; battery capacity degrades {BAT_DEG*100:.1f}%/yr. "
        f"Net system cost (gross price minus rebates) is treated as an upfront outlay at year 0."
    )
    doc.add_paragraph(
        f"Two payback metrics are shown. "
        f"The Nominal Payback is simply when cumulative electricity savings equal the upfront cost. "
        f"The Opp-Cost Payback discounts future savings at {OPPORTUNITY_RATE*100:.0f}%/yr — "
        f"it answers the question: when has the investment returned more than its cost in today's dollars? "
        f"NPV @ {OPPORTUNITY_RATE*100:.0f}% is the net present value of 25-yr savings at that rate — "
        f"positive means the investment is worthwhile at this hurdle rate."
    )
    pb_rows = []
    for pb in pbs:
        pb_rows.append({
            "Quote":         pb["label"][:28],
            "Tariff":        pb["tariff"],
            "Gross Cost":    f"${pb['cost']:,.0f}",
            "Fed Rebate":    f"${pb['fed']:,.0f}",
            "State Rebate":  f"${pb['state']:,.0f}",
            "Net Cost":      f"${pb['net']:,.0f}",
            "Payback":       f"{pb['pb_yr']} yr" if pb["pb_yr"] else ">25 yr",
            f"Opp-Cost Pb\n({OPPORTUNITY_RATE*100:.0f}%/yr)":
                             f"{pb['pb_yr_disc']} yr" if pb.get("pb_yr_disc") else ">25 yr",
            f"NPV @ {OPPORTUNITY_RATE*100:.0f}%":
                             f"${pb['npv']:,.0f}" if pb.get("npv") is not None else "n/a",
            "25-yr Saving":  f"${pb['total_save']:,.0f}",
            "ROI":           f"{pb['roi']:.0f}%",
            "Yr-1 Saving":   f"${pb['yr1_save']:,.0f}",
        })
    add_table(doc, pd.DataFrame(pb_rows))
    doc.add_paragraph(
        "Sensitivity: a 15% reduction in solar output (e.g. more shading) lengthens payback "
        "by ~2–3 years; tariff escalation at 5%/yr shortens it by ~1–2 years; every 10% off "
        "the system cost reduces breakeven proportionally."
    )

    # ── 8. Battery-Only Grid Arbitrage — omitted (no ARBIT_QUOTES defined) ──────
    if arbit_res:
        doc.add_heading("8. Battery-Only Grid Arbitrage (No Solar)", 1)
        doc.add_paragraph(
            f"Battery-only arbitrage on Midday Saver: charge cheaply during SOP (9am–3pm, "
            f"{MS_SOP*100:.4f}¢/kWh), discharge during peak (3–9pm, {MS_PEAK*100:.4f}¢/kWh). "
            f"Gross spread: {(MS_PEAK - MS_SOP)*100:.2f}¢/kWh; net after {BAT_RTE*100:.0f}% RTE: "
            f"~{(MS_PEAK - MS_SOP/BAT_RTE)*100:.2f}¢/kWh per kWh cycled."
        )
        arb = arbit_res[0]
        doc.add_paragraph(
            f"Modelled result for a {arb['bat_kwh']} kWh / {arb['inv_kw']} kW battery:"
        )
        p = doc.add_paragraph()
        bold_run(p, "Annual grid import: ");    p.add_run(f"{arb['annual_import_kwh']:,.0f} kWh\n")
        bold_run(p, "Net electricity cost: ");  p.add_run(f"${arb['net_cost']:,.0f}/year\n")
        bold_run(p, "Saving vs Midday Saver baseline: "); p.add_run(f"${base_ms - arb['net_cost']:,.0f}/year\n")
        bold_run(p, "% of peak load covered by battery: "); p.add_run(f"{arb['self_suf_pct']:.1f}%")
        if arbit_pbs:
            apb = arbit_pbs[0]
            doc.add_paragraph(
                f"Over {ANALYSIS_YEARS} years, net cost ${apb['net']:,.0f} after rebates. "
                f"Payback: {'%d years' % apb['pb_yr'] if apb['pb_yr'] else '>25 years'}. "
                f"Total saving: ${apb['total_save']:,.0f}.  ROI: {apb['roi']:.0f}%."
            )

    # ── 9. Zero-Net Sizing ────────────────────────────────────────────────────
    doc.add_heading("9. Zero-Net Sizing Analysis", 1)
    doc.add_paragraph(
        f"Zero-net = annual DEBS export credits ≥ total grid import cost + supply charges. "
        f"The table below shows the minimum solar array needed for each battery size on the "
        f"Midday Saver tariff at current DEBS rates. A larger battery requires less solar "
        f"(better self-consumption); a smaller battery needs more solar volume to compensate. "
        f"Size for well below zero-net in year 1 — DEBS rates decline {DEBS_DECL*100:.0f}%/yr."
    )
    if zero_net:
        zn_rows = []
        for r in zero_net:
            zn_rows.append({
                "Battery (kWh)":    f"{r['bat_kwh']:.0f}",
                "Min solar (kW)":   f"{r['min_solar_kw']:.1f}" if r["min_solar_kw"] is not None else ">35",
                "Net cost at min":  f"${r['net_cost']:,.0f}",
                "Export (kWh)":     f"{r['annual_export_kwh']:,.0f}",
                "Self-suf (%)":     f"{r['self_suf_pct']:.0f}%",
                "Achievable?":      "Yes" if r["min_solar_kw"] is not None else "No (>35 kW needed)",
            })
        add_table(doc, pd.DataFrame(zn_rows))
    doc.add_paragraph(
        "Note: single-phase properties in WA are typically limited to 5 kW inverter export. "
        "Confirm connection type and export approval with your installer before targeting "
        "a large-export zero-net design."
    )

    # ── 10. Solar Size Optimisation ───────────────────────────────────────────
    if sweep_df is not None and opt_kw is not None:
        doc.add_heading("10. Solar Size Optimisation (14 kWh Battery)", 1)
        doc.add_paragraph(
            f"Solar sizes from 3–20 kW tested in 0.5 kW steps, each paired with a 14 kWh battery "
            f"on Midday Saver. Recommended size: ~{opt_kw:.1f} kW — the 'knee' where self-sufficiency "
            f"gains flatten and additional panels mostly add export at low DEBS rates rather than "
            f"self-consumption. Larger arrays still save more in absolute terms but with longer "
            f"payback on the extra panel cost."
        )
        sweep_tbl = sweep_df[["solar_kw", "self_suf_pct", "bat_util_pct",
                               "days_full_pct", "export_kwh", "annual_savings"]].copy()
        sweep_tbl.columns = ["Solar (kW)", "Self-suf (%)", "Bat util (%)",
                              "Days full (%)", "Export (kWh)", "Saving ($)"]
        sweep_tbl = sweep_tbl.round(1)
        sweep_tbl["Saving ($)"] = sweep_tbl["Saving ($)"].apply(lambda x: f"${x:,.0f}")
        add_table(doc, sweep_tbl)
        doc.add_paragraph(
            "Columns: Self-suf = % of load met by solar+battery (not grid). "
            "Bat util = % of days where the battery reached at least 80% SoC. "
            "Days full = % of days the battery was fully charged at least once. "
            "Export = total annual kWh sold back to the grid. "
            "Saving = annual saving vs the Midday Saver baseline with no solar or battery."
        )

    # ── 11. Model Limitations and Caveats ────────────────────────────────────
    doc.add_heading("11. Model Limitations and Important Caveats", 1)
    for bullet in [
        "Synthetic solar: uses a calibrated sin(elevation) profile, not measured irradiance. "
        "Monthly production can deviate ±10–20% from the model; annual averages are more reliable.",
        "Single consumption year: 2024 usage pattern applied to all 25 years. Major lifestyle "
        "changes (EV, pool, household size) will alter actual savings.",
        f"Opportunity cost: the NPV and discounted payback columns use {OPPORTUNITY_RATE*100:.0f}%/yr "
        f"as the hurdle rate. Solar savings are lower-risk than market returns, so this "
        f"is a conservative discount rate.",
        "No maintenance costs: inverter replacement (~$1,500–3,000 once in 25 years) and "
        "cleaning costs are not included.",
        "Tariff uncertainty: Synergy may restructure tariffs at any time over the 25-year horizon.",
        "Export limits: single-phase properties are capped at 5 kW inverter export by Western Power.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    # ── 12. Charts ────────────────────────────────────────────────────────────
    doc.add_heading("12. Charts", 1)
    chart_files = [
        ("17_monthly_bill_all_scenarios.png",
         "Fig 1 — Predicted Monthly Electricity Bill — All Scenarios\n"
         "Net monthly Synergy bill for each scenario vs the current baseline (thick black line). "
         "Solid = A1 Flat; dashed = Midday Saver. "
         "Summer dip is largest because panels generate most electricity in Dec–Feb."),
        ("00_consumption_and_shading.png",
         "Fig 2 — Your Actual 2024 Consumption & Site Shading Model\n"
         "Left: half-hourly consumption heatmap (columns = months, rows = time of day). "
         "Darker red = more electricity used. The red band marks the 3–9 pm peak window. "
         "Right: seasonal shading model — fraction of maximum solar output reaching the panels."),
        ("01_average_daily_profile.png",
         "Fig 3 — Average Daily Energy Profile (A1 Flat top, Midday Saver bottom)\n"
         "Average half-hourly energy flows for the first solar+battery quote. Yellow = solar "
         "self-consumption; green = battery discharge; red = grid import; dashed yellow = total "
         "solar generation. Green/red shading marks the SOP and peak windows."),
        ("04_monthly_costs.png",
         "Fig 4 — Monthly Electricity Costs by Tariff\n"
         "Net monthly cost for the first quote on each tariff (blue = A1, green = Midday Saver) "
         "vs the A1 Flat baseline (black step line). Bars below zero mean DEBS credits exceed imports."),
        ("14_monthly_cost_verification.png",
         "Fig 5 — VERIFICATION: Modelled Baseline vs All Scenarios\n"
         "CHECK THIS FIRST. The solid line is the modelled grid-only baseline — compare it to "
         "your actual Synergy bills month by month. A persistent offset likely means a separately "
         "metered circuit (e.g. hot-water) is missing from the smart meter export."),
        ("06_payback_analysis.png",
         f"Fig 6 — Solar+Battery Cumulative Cash-Flow ({ANALYSIS_YEARS}-Year Horizon)\n"
         f"Each colour = one quote. Solid lines = A1 Flat tariff; dashed lines = Midday Saver. "
         f"Lines start at −net cost at year 0 and rise with cumulative annual savings. "
         f"The break-even year (when the line crosses zero) is labelled in the legend."),
        ("08_solar_sweep.png",
         "Fig 7 — Solar Size Optimisation (14 kWh Battery, Midday Saver)\n"
         f"Self-sufficiency (orange), battery utilisation (green), and annual export (red dotted) "
         f"vs solar array size 3–20 kW. The dashed vertical line at {opt_kw:.1f} kW marks the "
         "'knee' — beyond this, additional panels mostly add export rather than self-consumption."),
        ("13_monthly_savings.png",
         "Fig 8 — Monthly Savings vs No-Solar Baseline\n"
         "Monthly saving vs the same-tariff baseline for every quote. Solid = A1 Flat; "
         "hatched = Midday Saver. Positive = money saved; higher bars in summer reflect "
         "peak solar production."),
        ("16_zero_net_sizing.png",
         "Fig 10 — Zero-Net Sizing: Minimum Solar Array per Battery Size\n"
         "Minimum solar array (kW) needed for zero net annual cost, for each battery size. "
         "Green = achievable within 35 kW; red = not achievable. Larger batteries require "
         "smaller arrays."),
        ("09_summer_day.png",
         "Fig 11 — Representative Summer Day (December–February)\n"
         "Half-hourly energy flows for a typical summer day — first solar+battery quote. "
         "Large midday solar peak charges the battery and exports surplus; battery then "
         "covers the evening peak. Bottom panel shows battery state-of-charge."),
        ("10_autumn_day.png",
         "Fig 12 — Representative Autumn Day (March–May)\n"
         "Moderate solar output as days shorten. Battery typically charges to 70–90% capacity "
         "and still covers a meaningful share of the evening peak demand."),
        ("11_winter_day.png",
         "Fig 13 — Representative Winter Day (June–August)\n"
         "Shorter days, lower sun arc — battery charges to ~50–70% on a typical winter day. "
         "Evening grid import rises, but on Midday Saver most daytime import is at the cheap "
         "SOP rate, limiting the cost impact."),
        ("12_spring_day.png",
         "Fig 14 — Representative Spring Day (September–November)\n"
         "Solar output and battery cycles recover rapidly through spring. By October the "
         "system approaches summer performance, with full daily charge/discharge cycles "
         "and low evening grid imports."),
    ]
    for fname, caption in chart_files:
        fpath = out_dir / fname
        if fpath.exists():
            p = doc.add_paragraph()
            p.add_run(caption).italic = True
            doc.add_picture(str(fpath), width=Inches(6.0))
            doc.add_paragraph()

    # ── 13. Assumptions ───────────────────────────────────────────────────────
    doc.add_heading("13. Key Assumptions & Methodology", 1)
    doc.add_paragraph(
        "All assumptions can be modified in the configuration section at the top of the "
        "Python script (solar_battery_analyser.py). If you receive updated installer quotes, "
        "a shading assessment, or new tariff information, update the relevant constants and "
        "re-run the analysis."
    )
    for text in [
        f"Tariff rates — Synergy WA, effective 1 July 2025. "
        f"A1 Flat: {A1_RATE*100:.4f}¢/kWh usage rate + {A1_SUPPLY*100:.4f}¢/day supply charge (fixed daily). "
        f"Midday Saver: {MS_SOP*100:.4f}¢/kWh super off-peak (9am–3pm), "
        f"{MS_PEAK*100:.4f}¢/kWh peak (3–9pm), {MS_OP*100:.4f}¢/kWh off-peak (9pm–9am). "
        f"Midday Saver supply charge: {MS_SUPPLY*100:.4f}¢/day.",
        f"DEBS export rates — Distributed Energy Buyback Scheme: {DEBS_PEAK*100:.0f}¢/kWh "
        f"during the peak window (3–9 pm), {DEBS_OP*100:.0f}¢/kWh at all other times. "
        f"DEBS rates decline {DEBS_DECL*100:.0f}% per year over the analysis period, "
        f"reflecting the ongoing reduction in the value of grid-connected solar export.",
        f"Tariff escalation — {TARIFF_ESC*100:.1f}% per annum applied to all import charges "
        f"and supply charges throughout the {ANALYSIS_YEARS}-year horizon. This is a conservative "
        f"base case; Synergy has historically increased tariffs by 3–6%/year. Sensitivity: "
        f"at 5%/year escalation most payback periods shorten by approximately 1–2 years.",
        f"Battery parameters — round-trip efficiency {BAT_RTE*100:.0f}% (√ applied symmetrically "
        f"to charge and discharge halves); depth of discharge {BAT_DOD*100:.0f}% of nameplate; "
        f"annual capacity degradation {BAT_DEG*100:.1f}%/yr. Inverter dispatch limited to rated "
        f"inverter size (kW) per quote.",
        f"Solar modelling — synthetic bell-curve daily profile scaled by panel array size (kW DC), "
        f"system efficiency {SYS_EFF*100:.0f}%, seasonal Peak Sun Hours (4 hrs winter – 6 hrs summer), "
        f"seasonal shading factors, and diurnal sunrise/sunset taper. "
        f"Panel degradation {SOL_DEG*100:.2f}%/yr applied annually.",
        f"Shading model — seasonal factors: summer {SHADE_SUMMER*100:.0f}%, "
        f"autumn {SHADE_AUTUMN*100:.0f}%, winter {SHADE_WINTER*100:.0f}%, "
        f"spring {SHADE_SPRING*100:.0f}%. Plus a diurnal taper to zero at sunrise/sunset. "
        f"These factors produce approximately 40% of theoretical unshaded annual output — "
        f"adjust upward for an unshaded north-facing roof.",
        f"Federal battery rebate — Cheaper Home Batteries Program (effective 1 May 2026): "
        f"tiered at $252/kWh (0–14 kWh), $151/kWh (14–28 kWh), $38/kWh (28–50 kWh), "
        f"capped at 40% of gross installed system cost. Applied to battery kWh only.",
        f"State battery rebate — Synergy WA Residential Battery Scheme: "
        f"${STATE_REBATE_FLAT:,.0f} flat payment, requires VPP enrolment. "
        f"All battery quotes assumed to be VPP-eligible.",
        f"STC solar rebate — assumed already reflected in the installer's gross quoted price. "
        f"Not modelled separately in this analysis.",
        f"Analysis horizon — {ANALYSIS_YEARS} years. Nominal payback figures are in nominal AUD "
        f"(not inflation-adjusted). Discounted payback and NPV use {OPPORTUNITY_RATE*100:.0f}%/yr "
        f"as the hurdle rate. Assumes cash purchase. No maintenance or replacement costs included.",
        "Data source — half-hourly smart meter consumption data (kWh per 30-minute interval) "
        "exported from Synergy online account. Controlled-load and separately metered circuits "
        "may not be included; check against your actual Synergy bills.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    # ── Conclusion & Recommendation ───────────────────────────────────────────
    doc.add_heading("Conclusion & Recommendation", 1)

    # All non-base-case quotes (solar+battery and battery-only arbitrage)
    conclusion_pbs = [pb for pb in pbs if pb["net"] > 0]
    if arbit_pbs:
        conclusion_pbs += list(arbit_pbs)

    def _best_at(yr):
        """Quote with highest cumulative cashflow at a given year."""
        return max(conclusion_pbs, key=lambda pb: pb["cum"][yr])

    for horizon in [10, 20]:
        doc.add_heading(f"{horizon}-Year Outlook", 2)
        best     = _best_at(horizon)
        best_cum = best["cum"][horizon]
        doc.add_paragraph(
            f"Recommended: '{best['label']}' on the {best['tariff']} tariff. "
            f"By year {horizon} this system delivers a cumulative saving of ${best_cum:,.0f}. "
            f"Nominal payback: {best['pb_yr']} yr  ·  NPV @ {OPPORTUNITY_RATE*100:.0f}%: "
            f"${best['npv']:,.0f}."
        )

    # ── Overall summary ───────────────────────────────────────────────────────
    doc.add_heading("Overall Recommendation", 2)

    best_25  = max(conclusion_pbs, key=lambda pb: pb["cum"][ANALYSIS_YEARS])
    best_npv = max(conclusion_pbs, key=lambda pb: pb.get("npv", 0))

    rec_text = (
        f"Based on financial return, the standout option is "
        f"'{best_25['label']}' on the {best_25['tariff']} tariff: "
        f"nominal payback of {best_25['pb_yr']} years, "
        f"25-year cumulative saving of ${best_25['total_save']:,.0f}, "
        f"and NPV of ${best_25['npv']:,.0f} at {OPPORTUNITY_RATE*100:.0f}%/yr."
    )
    if best_npv["label"] != best_25["label"] or best_npv["tariff"] != best_25["tariff"]:
        rec_text += (
            f" Best NPV per dollar invested: '{best_npv['label']}' on "
            f"{best_npv['tariff']} (${best_npv['npv']:,.0f})."
        )
    doc.add_paragraph(rec_text)

    doc.add_paragraph(
        f"Important: the above is a financial comparison only. Non-financial benefits "
        f"of solar+battery include protection against electricity tariff increases "
        f"beyond the modelled {TARIFF_ESC*100:.1f}%/yr (historically 3–6%/yr), "
        f"reduced carbon footprint, potential grid-outage resilience (if your battery "
        f"supports backup mode), and reduced exposure to future energy market uncertainty. "
        f"These factors may justify the investment even where the pure financial case "
        f"is marginal."
    )

    # ── Sizing Optimisation — Self-Consumption vs Export ──────────────────────
    doc.add_heading("Sizing Optimisation — Self-Consumption vs Export", 1)
    sa = analyse_sizing(pbs, opt_kw=opt_kw)
    avg_fit_c = (0.25 * DEBS_PEAK + 0.75 * DEBS_OP) * 100   # ¢/kWh

    doc.add_paragraph(
        f"A key driver of financial return under your site conditions is self-consumption: "
        f"the fraction of solar generation consumed on-site (or stored in the battery) "
        f"rather than exported. Your DEBS export tariff averages approximately "
        f"{avg_fit_c:.1f}¢/kWh (10¢ peak, 2¢ off-peak), compared with an import rate "
        f"of {A1_RATE*100:.1f}¢/kWh — meaning exported energy earns only "
        f"{avg_fit_c/(A1_RATE*100)*100:.0f}% of its import value. "
        f"With seasonal shading reducing output by 40–60%, oversizing the array "
        f"beyond what the battery and on-site load can absorb sends excess generation "
        f"to the grid at low rates, diluting returns."
    )

    if sa and sa["rows"]:
        # Table: each Midday Saver solar quote ranked by NPV
        tbl_data_rows = []
        for row in sa["rows"]:
            export_str = (f"{row['export_ratio']:.0f}%"
                          if row["solar_kwh"] > 0 else "—")
            tbl_data_rows.append({
                "System": row["label"],
                "Net Cost": f"${row['net']:,.0f}",
                "Self-suf": f"{row['self_suf_pct']:.0f}%",
                "Export\n(kWh/yr)": f"{row['export_kwh']:,.0f}",
                "Export\nRatio": export_str,
                "FiT Loss\n($/yr)": f"${row['annual_fit_loss']:,.0f}",
                "NPV\n(25 yr)": f"${row['npv']:,.0f}",
                "NPV per\n$1k invested": f"${row['npv_per_k']:,.0f}",
            })
        tbl_df = pd.DataFrame(tbl_data_rows)
        add_table(doc, tbl_df)
        doc.add_paragraph(
            "Export Ratio = exported kWh ÷ total solar generated. "
            "FiT Loss = value destroyed annually by exporting at 2–10¢ instead of self-consuming at 32¢. "
            "NPV = cumulative 25-year discounted savings, net of upfront cost. "
            "NPV per $1k = NPV divided by net capex per $1,000 invested (capital efficiency)."
        ).runs[0].font.size = Pt(7)

        bn = sa["best_npv"]
        be = sa["best_eff"]
        doc.add_paragraph(
            f"Highest absolute NPV: '{bn['label']}' — ${bn['npv']:,.0f} over 25 years, "
            f"with {bn['self_suf_pct']:.0f}% self-sufficiency and "
            f"{bn['export_ratio']:.0f}% of solar exported."
        )
        if be["label"] != bn["label"]:
            doc.add_paragraph(
                f"Best capital efficiency: '{be['label']}' — ${be['npv_per_k']:,.0f} NPV per "
                f"$1,000 invested, with {be['self_suf_pct']:.0f}% self-sufficiency and "
                f"{be['export_ratio']:.0f}% of solar exported. "
                f"This option delivers more return per dollar of capex, which is relevant "
                f"if the higher-cost options are competing on the margin."
            )
        else:
            doc.add_paragraph(
                f"This option is also the most capital-efficient at "
                f"${be['npv_per_k']:,.0f} NPV per $1,000 invested."
            )

        # Flag high-export-ratio options
        high_export = [r for r in sa["rows"] if r["export_ratio"] > 40 and r["solar_kwh"] > 0]
        if high_export:
            names = ", ".join(f"'{r['label']}'" for r in high_export)
            doc.add_paragraph(
                f"Caution — potential oversizing: {names} "
                f"export more than 40% of their solar generation. Given the low DEBS FiT rate "
                f"and your site shading, the marginal panels on these larger arrays are unlikely "
                f"to recover their cost within the analysis horizon. "
                f"A smaller, well-matched array with adequate battery storage will typically "
                f"deliver better capital efficiency."
            )

        if opt_kw:
            doc.add_paragraph(
                f"The solar self-sufficiency sweep (Section 9, 14 kWh battery) shows "
                f"diminishing returns beyond {opt_kw:.1f} kW — additional panels beyond "
                f"this point primarily add export rather than on-site value. "
                f"Under your shading conditions, sizing close to this knee maximises "
                f"return per dollar of solar capex."
            )

    path = out_dir / "solar_analysis_report.docx"
    doc.save(str(path))
    print(f"  Word report  → {path.resolve()}")
    return path


# ─── PDF REPORT ───────────────────────────────────────────────────────────────
def write_pdf_report(all_res, pbs, raw_df, opt_kw, out_dir,
                     arbit_res=None, arbit_pbs=None, zero_net=None):
    """Multi-page PDF: title page, financial tables, all charts with captions."""
    from matplotlib.backends.backend_pdf import PdfPages
    import datetime as dt

    # Matplotlib treats '$' as LaTeX math delimiter — escape it in all text strings.
    def _t(s): return s.replace('$', r'\$')

    # ── Reusable table style: blue header row, alternating grey rows, light border
    def _style_tbl(tbl):
        for (r, _), cell in tbl.get_celld().items():
            if r == 0:   # header row
                cell.set_facecolor("#2c5f8a")
                cell.set_text_props(color="white", fontweight="bold")
            elif r % 2 == 0:   # alternating shading for readability
                cell.set_facecolor("#eef3f8")
            cell.set_edgecolor("#cccccc")

    base_a1  = baseline(raw_df, "A1 Flat")
    base_ms  = baseline(raw_df, "Midday Saver")
    total    = raw_df["consumption_kwh"].sum()
    n_days   = int(raw_df["date"].nunique()) if "date" in raw_df.columns else 365

    path = out_dir / "solar_analysis_report.pdf"

    with PdfPages(str(path)) as pdf:

        # ── Page 1: Title & consumption summary ──────────────────────────────
        fig = plt.figure(figsize=(11, 8.5))
        fig.patch.set_facecolor("#f5f5f5")
        ax = fig.add_axes([0, 0, 1, 1]); ax.axis("off")

        ax.text(0.5, 0.90, "Solar & Battery Economic Analysis",
                ha="center", va="top", fontsize=26, fontweight="bold",
                transform=ax.transAxes)
        ax.text(0.5, 0.83,
                f"Synergy WA (Perth / SWIS)  ·  Generated {dt.date.today().strftime('%d %B %Y')}",
                ha="center", va="top", fontsize=13, color="#555555",
                transform=ax.transAxes)

        # Consumption summary box
        summary_lines = [
            ("Annual consumption",      f"{total:,.0f} kWh"),
            ("Average daily",           f"{total/n_days:.1f} kWh/day"),
            ("Data days",               f"{n_days}"),
            ("",                        ""),
            ("Baseline — A1 Flat",      f"${base_a1:,.2f}/yr  (${base_a1/12:,.2f}/mo)"),
            ("Baseline — Midday Saver", f"${base_ms:,.2f}/yr  (${base_ms/12:,.2f}/mo)"),
            ("",                        ""),
            ("A1 usage rate",           f"{A1_RATE*100:.4f} ¢/kWh"),
            ("A1 supply charge",        f"{A1_SUPPLY*100:.4f} ¢/day"),
            ("MS peak rate (3–9 pm)",   f"{MS_PEAK*100:.4f} ¢/kWh"),
            ("MS shoulder (9 am–3 pm)", f"{MS_SOP*100:.4f} ¢/kWh"),
            ("MS off-peak",             f"{MS_OP*100:.4f} ¢/kWh"),
            ("DEBS export (peak)",      f"{DEBS_PEAK*100:.0f} ¢/kWh"),
            ("DEBS export (other)",     f"{DEBS_OP*100:.0f} ¢/kWh"),
        ]
        y0 = 0.73
        for label, value in summary_lines:
            if not label:
                y0 -= 0.012
                continue
            ax.text(0.22, y0, label + ":", ha="right", fontsize=10,
                    transform=ax.transAxes, color="#333333")
            ax.text(0.24, y0, _t(value), ha="left", fontsize=10, fontweight="bold",
                    transform=ax.transAxes)
            y0 -= 0.038

        ax.text(0.5, 0.06,
                "This report was generated automatically from your Synergy smart meter data and "
                "synthetic solar profiles calibrated to Perth conditions.\n"
                "All figures are model estimates — actual performance will depend on weather, shading, "
                "installation quality, and future tariff movements.\n"
                "Read this report alongside installer quotes and a professional site shading assessment. "
                "Verify the baseline against your actual Synergy bills using Fig 5.",
                ha="center", va="bottom", fontsize=8.5, color="#777777",
                transform=ax.transAxes)
        pdf.savefig(fig, bbox_inches="tight"); plt.close(fig)

        # ── Page 2: Year-1 results table ─────────────────────────────────────
        fig, ax = plt.subplots(figsize=(11, 8.5))
        ax.axis("off")
        fig.patch.set_facecolor("white")
        ax.set_title("Year-1 Results by Quote & Tariff",
                     fontsize=16, fontweight="bold", pad=20)

        yr1_rows = []
        for r in list(all_res) + (list(arbit_res) if arbit_res else []):
            base = base_a1 if r["tariff"] == "A1 Flat" else base_ms
            yr1_rows.append([
                r["label"][:28],
                r["tariff"],
                f"{r['annual_solar_kwh']:,.0f}",
                f"{r['annual_import_kwh']:,.0f}",
                f"{r['annual_export_kwh']:,.0f}",
                f"${r['net_cost']:,.0f}",
                f"${base - r['net_cost']:,.0f}",
                f"{r['self_suf_pct']:.0f}%",
            ])
        col_labels = ["Quote", "Tariff", "Solar\n(kWh)", "Import\n(kWh)",
                      "Export\n(kWh)", "Net\nCost", "Saving\nvs base", "Self-\nsuf"]
        tbl = ax.table(cellText=yr1_rows, colLabels=col_labels,
                       loc="center", cellLoc="center")
        tbl.auto_set_font_size(False); tbl.set_fontsize(8)
        tbl.scale(1, 1.4)
        _style_tbl(tbl)
        ax.text(0.5, 0.03,
                "Year-1 = first full year of operation before tariff escalation or equipment degradation.\n"
                "Net Cost = (import kWh × tariff rate) + (supply charge × 365 days) − (export kWh × DEBS rate).  "
                "Saving = same-tariff baseline minus Net Cost.  "
                "Self-suf = % of total load met by solar self-consumption + battery discharge (not from grid).\n"
                "Savings grow in later years as tariff escalation applies — see the 25-year financial table.",
                ha="center", va="bottom", fontsize=7.5, color="#555555",
                transform=ax.transAxes)
        pdf.savefig(fig, bbox_inches="tight"); plt.close(fig)

        # ── Page 3: Financial / payback table ────────────────────────────────
        fig, ax = plt.subplots(figsize=(11, 8.5))
        ax.axis("off")
        fig.patch.set_facecolor("white")
        ax.set_title(f"Financial Summary — {ANALYSIS_YEARS}-Year Horizon",
                     fontsize=14, fontweight="bold", pad=20)

        all_pbs = list(pbs) + (list(arbit_pbs) if arbit_pbs else [])
        pb_rows = []
        for pb in all_pbs:
            pb_rows.append([
                pb["label"][:24],
                pb["tariff"],
                f"${pb['cost']:,.0f}",
                f"${pb['net']:,.0f}",
                f"{pb['pb_yr']} yr" if pb["pb_yr"] else ">25 yr",
                f"{pb['pb_yr_disc']} yr" if pb.get("pb_yr_disc") else ">25 yr",
                f"${pb.get('npv', 0):,.0f}",
                f"${pb['total_save']:,.0f}",
                f"{pb['roi']:.0f}%",
            ])
        pb_cols = ["Quote", "Tariff", "Gross\nCost", "Net\nCost",
                   "Nominal\nPayback", f"Opp-Cost\nPayback\n({OPPORTUNITY_RATE*100:.0f}%)",
                   f"NPV\n@ {OPPORTUNITY_RATE*100:.0f}%", "25-yr\nSaving", "ROI"]
        tbl2 = ax.table(cellText=pb_rows, colLabels=pb_cols,
                        loc="center", cellLoc="center")
        tbl2.auto_set_font_size(False); tbl2.set_fontsize(8)
        tbl2.scale(1, 1.3)
        _style_tbl(tbl2)
        ax.text(0.5, 0.03,
                _t(f"Net Cost = gross price minus federal rebate (tiered $252/$151/$38/kWh, capped 40%) and "
                   f"state rebate (${STATE_REBATE_FLAT:,.0f} flat, VPP required).  "
                   f"Nominal Payback = year cumulative savings first equal net cost (no discounting).  "
                   f"Opp-Cost Payback = year discounted savings (at {OPPORTUNITY_RATE*100:.0f}%/yr) first exceed net cost.\n"
                   f"NPV = net present value of 25-yr savings discounted at {OPPORTUNITY_RATE*100:.0f}%/yr (positive = worthwhile).  "
                   f"ROI = total 25-yr saving / net cost.\n"
                   f"Assumptions: tariff escalation {TARIFF_ESC*100:.1f}%/yr; DEBS declining {DEBS_DECL*100:.0f}%/yr; "
                   f"panel degradation {SOL_DEG*100:.2f}%/yr; battery degradation {BAT_DEG*100:.1f}%/yr."),
                ha="center", va="bottom", fontsize=7, color="#555555",
                transform=ax.transAxes)
        pdf.savefig(fig, bbox_inches="tight"); plt.close(fig)

        # ── Page 4: Conclusion & Recommendation ──────────────────────────────
        conclusion_pbs = [pb for pb in pbs if pb["net"] > 0]
        if arbit_pbs:
            conclusion_pbs += list(arbit_pbs)

        def _best_at_pdf(yr):
            return max(conclusion_pbs, key=lambda pb: pb["cum"][yr])

        fig, ax = plt.subplots(figsize=(11, 8.5))
        ax.axis("off"); fig.patch.set_facecolor("white")
        ax.set_title("Conclusion & Recommendation",
                     fontsize=16, fontweight="bold", pad=20)

        lines = []
        for horizon in [10, 20]:
            best     = _best_at_pdf(horizon)
            best_cum = best["cum"][horizon]
            s_cum    = _t(f"${best_cum:,.0f}")

            lines.append(f"── {horizon}-Year Outlook ──")
            lines.append(f"Recommended: '{best['label']}' on {best['tariff']}.")
            lines.append(
                f"  Cumulative saving at yr {horizon}: {s_cum}  "
                f"· Nominal payback: {best['pb_yr']} yr  "
                f"· NPV @ {OPPORTUNITY_RATE*100:.0f}%: ${best['npv']:,.0f}")
            lines.append("")

        # Overall
        best_25  = max(conclusion_pbs, key=lambda pb: pb["cum"][ANALYSIS_YEARS])
        best_npv = max(conclusion_pbs, key=lambda pb: pb.get("npv", 0))

        lines.append("── Overall Recommendation ──")
        s_best_save = _t(f"${best_25['total_save']:,.0f}")
        lines.append(f"Financially strongest option: '{best_25['label']}' on {best_25['tariff']}.")
        lines.append(
            f"  Nominal payback: {best_25['pb_yr']} yr  ·  "
            f"25-yr saving: {s_best_save}  ·  "
            f"NPV: ${best_25['npv']:,.0f}")
        if best_npv["label"] != best_25["label"] or best_npv["tariff"] != best_25["tariff"]:
            s_npv_save = _t(f"${best_npv['total_save']:,.0f}")
            lines.append(
                f"  Best NPV per dollar: '{best_npv['label']}' on {best_npv['tariff']} "
                f"(NPV ${best_npv['npv']:,.0f}, saving {s_npv_save}).")
        lines.append("")
        lines.append(
            f"Note: non-financial benefits (tariff-rise protection beyond "
            f"{TARIFF_ESC*100:.1f}%/yr, carbon footprint, grid resilience) may justify")
        lines.append("  the investment even where the pure financial case is marginal.")

        y0 = 0.88
        for line in lines:
            bold = line.startswith("──")
            ax.text(0.05, y0, line, ha="left", va="top", fontsize=10,
                    fontweight="bold" if bold else "normal",
                    color="#1a1a1a" if bold else "#333333",
                    transform=ax.transAxes)
            y0 -= 0.052 if bold else 0.042
            if y0 < 0.05:
                break

        pdf.savefig(fig, bbox_inches="tight"); plt.close(fig)

        # ── Page 5: Sizing Optimisation — Self-Consumption vs Export ─────────
        sa = analyse_sizing(pbs, opt_kw=opt_kw)
        if sa and sa["rows"]:
            fig, ax = plt.subplots(figsize=(11, 8.5))
            ax.axis("off"); fig.patch.set_facecolor("white")
            ax.set_title("Sizing Optimisation — Self-Consumption vs Export",
                         fontsize=14, fontweight="bold", pad=20)

            avg_fit_c = (0.25 * DEBS_PEAK + 0.75 * DEBS_OP) * 100
            fit_pen_c = (A1_RATE - (0.25 * DEBS_PEAK + 0.75 * DEBS_OP)) * 100

            # Header text
            intro_lines = [
                f"DEBS export rate ≈ {avg_fit_c:.1f}¢/kWh avg (10¢ peak, 2¢ off-peak)  vs  "
                f"import rate {A1_RATE*100:.1f}¢/kWh.",
                f"Each exported kWh destroys {fit_pen_c:.1f}¢ of potential value — "
                f"self-consumption is critical.",
                f"Shading factors: Summer {SHADE_SUMMER*100:.0f}%  Autumn {SHADE_AUTUMN*100:.0f}%  "
                f"Winter {SHADE_WINTER*100:.0f}%  Spring {SHADE_SPRING*100:.0f}%",
                "",
            ]
            y0 = 0.88
            for ln in intro_lines:
                ax.text(0.03, y0, ln, ha="left", va="top", fontsize=9,
                        color="#444444", transform=ax.transAxes)
                y0 -= 0.038

            # Table of quotes ranked by NPV
            col_labels = ["System", "Net\nCost", "Self-\nsuf%", "Export\nkWh/yr",
                          "Export\nRatio", "FiT Loss\n$/yr", "NPV\n25yr", "NPV/\n$1k"]
            tbl_rows = []
            for row in sa["rows"]:
                exp_str = f"{row['export_ratio']:.0f}%" if row["solar_kwh"] > 0 else "—"
                tbl_rows.append([
                    row["label"][:30],
                    f"${row['net']:,.0f}",
                    f"{row['self_suf_pct']:.0f}%",
                    f"{row['export_kwh']:,.0f}",
                    exp_str,
                    f"${row['annual_fit_loss']:,.0f}",
                    f"${row['npv']:,.0f}",
                    f"${row['npv_per_k']:,.0f}",
                ])

            tbl = ax.table(cellText=tbl_rows, colLabels=col_labels,
                           loc="upper center", cellLoc="center",
                           bbox=[0.0, y0 - 0.02 - len(tbl_rows)*0.07, 1.0,
                                 0.04 + (len(tbl_rows)+1)*0.07])
            tbl.auto_set_font_size(False); tbl.set_fontsize(8)
            _style_tbl(tbl)

            # Highlight best-NPV row green, best-efficiency row blue
            bn_label = sa["best_npv"]["label"]
            be_label = sa["best_eff"]["label"]
            for (r, c), cell in tbl.get_celld().items():
                if r == 0:
                    continue
                row_label = sa["rows"][r-1]["label"] if r-1 < len(sa["rows"]) else ""
                if row_label == bn_label:
                    cell.set_facecolor("#d4edda")   # green — best NPV
                elif row_label == be_label:
                    cell.set_facecolor("#cce5ff")   # blue — best efficiency

            # Summary text below table
            y_txt = y0 - 0.05 - (len(tbl_rows)+1) * 0.07
            bn = sa["best_npv"]; be = sa["best_eff"]
            summary_lines = [
                "",
                f"Best NPV (green): '{bn['label']}'  —  ${bn['npv']:,.0f} over 25 yrs  |  "
                f"{bn['export_ratio']:.0f}% exported  |  {bn['self_suf_pct']:.0f}% self-sufficient",
            ]
            if be["label"] != bn["label"]:
                summary_lines.append(
                    f"Best capital efficiency (blue): '{be['label']}'  —  "
                    f"${be['npv_per_k']:,.0f} NPV per $1k invested  |  "
                    f"{be['export_ratio']:.0f}% exported"
                )
            high_exp = [r for r in sa["rows"] if r["export_ratio"] > 40 and r["solar_kwh"] > 0]
            if high_exp:
                names = "; ".join(r["label"][:25] for r in high_exp)
                summary_lines.append(
                    f"Potential oversizing (>40% export): {names}"
                )
                summary_lines.append(
                    "  → With low FiT and site shading, the marginal panels on larger arrays "
                    "are unlikely to recover their cost."
                )
            if opt_kw:
                summary_lines.append(
                    f"Solar sweep knee (14 kWh battery): {opt_kw:.1f} kW  —  "
                    "additional panels beyond this mostly add export, not savings."
                )

            for ln in summary_lines:
                bold = ln.startswith("Best") or ln.startswith("Potential") or ln.startswith("Solar")
                ax.text(0.03, y_txt, ln, ha="left", va="top", fontsize=9,
                        fontweight="bold" if bold else "normal",
                        color="#1a1a1a" if bold else "#444444",
                        transform=ax.transAxes)
                y_txt -= 0.042
                if y_txt < 0.02:
                    break

            pdf.savefig(fig, bbox_inches="tight"); plt.close(fig)

        # ── Page 6: Zero-net sizing table (if available) ─────────────────────
        if zero_net:
            fig, ax = plt.subplots(figsize=(11, 8.5))
            ax.axis("off")
            fig.patch.set_facecolor("white")
            ax.set_title("Zero-Net Sizing — Minimum Solar for $0 Annual Electricity Cost",
                         fontsize=14, fontweight="bold", pad=20)
            zn_rows = [[
                f"{r['bat_kwh']:.0f} kWh",
                f"{r['min_solar_kw']:.1f} kW" if r["min_solar_kw"] is not None else ">35 kW",
                f"${r['net_cost']:,.0f}",
                f"{r['annual_export_kwh']:,.0f} kWh",
                f"{r['self_suf_pct']:.0f}%",
                "YES" if r["min_solar_kw"] is not None else "NO",
            ] for r in zero_net]
            zn_cols = ["Battery", "Min Solar", "Net Cost\nat min solar",
                       "Annual Export", "Self-suf", "Zero-net\nachievable?"]
            tbl3 = ax.table(cellText=zn_rows, colLabels=zn_cols,
                            loc="center", cellLoc="center")
            tbl3.auto_set_font_size(False); tbl3.set_fontsize(9)
            tbl3.scale(1, 2.0)
            _style_tbl(tbl3)
            for (r, c), cell in tbl3.get_celld().items():
                if c == 5 and r > 0:
                    val = zn_rows[r - 1][5]
                    cell.set_facecolor("#d4edda" if val == "YES" else "#f8d7da")
            ax.text(0.5, 0.12,
                    _t("Zero-net = annual DEBS export credits >= annual import charges + supply charges — "
                       "your net Synergy bill is $0 or better.  Tariff: Midday Saver.  "
                       "Uses year-0 DEBS rates (current); DEBS declines 5%/yr making zero-net harder over time.\n"
                       "A larger battery increases self-consumption and can shift late-afternoon export to the "
                       "higher-value peak DEBS window — reducing the solar array size needed to reach zero-net.\n"
                       "Check with your installer regarding Western Power single-phase export limits (5 kW) "
                       "before designing a system specifically for maximum export."),
                    ha="center", va="bottom", fontsize=8, color="#555555",
                    transform=ax.transAxes)
            pdf.savefig(fig, bbox_inches="tight"); plt.close(fig)

        # ── Remaining pages: each chart with a caption ───────────────────────
        chart_pages = [
            ("17_monthly_bill_all_scenarios.png",
             "Fig 1 — Predicted Monthly Electricity Bill — All Scenarios",
             "Each line shows what your net monthly Synergy bill would look like under a given scenario. "
             "Thick black = current baseline (grid only, A1 Flat). Solid coloured = A1 Flat solar+battery; "
             "dashed = Midday Saver solar+battery. "
             "Lower lines = cheaper bill. Summer (Dec–Feb) shows the steepest savings because long days "
             "and high sun angles maximise solar output, often eliminating daytime grid imports entirely. "
             "Winter (Jun–Aug) savings are smaller due to shorter days and the sun tracking a low arc across the sky. "
             "Use this chart to estimate the monthly cash-flow improvement you would experience."),
            ("00_consumption_and_shading.png",
             "Fig 2 — Your Actual 2024 Consumption & Site Shading Model",
             "Left: heatmap of every half-hour of your 2024 electricity consumption. "
             "Horizontal axis = month (Jan–Dec); vertical axis = time of day (midnight top and bottom, noon middle). "
             "Darker red = more consumption in that slot. This reveals your daily load shape across seasons — "
             "summer afternoon AC load, winter morning and evening heating, and overnight base load. "
             "The red horizontal band marks the expensive 3–9 pm peak window on Midday Saver. "
             "Right: shading model — fraction of theoretical solar output reaching your panels each half-hour. "
             "Lighter = more sun reaching panels; darker = more shading. "
             "Summer (left) is mostly light (sun high); winter (right) is darker (sun low, longer shadows)."),
            ("01_average_daily_profile.png",
             "Fig 3 — Average Daily Energy Profile (A1 Flat top; Midday Saver bottom)",
             "Stacked areas show how load is met across a typical day for the first solar+battery quote. "
             "Yellow = solar self-consumed directly as generated. Green = battery discharging to cover load. "
             "Red = grid import. Black line = your actual average load. "
             "Dashed yellow = total solar output (includes portion exported to grid beyond the stacked area). "
             "Green background band = cheap SOP window (9am–3pm on Midday Saver). "
             "Red background band = expensive peak window (3–9pm). "
             "The chart shows how solar displaces grid imports during daylight hours, and how the battery "
             "covers the evening peak load, minimising the expensive red area during the peak window."),
            ("04_monthly_costs.png",
             "Fig 4 — Monthly Electricity Costs by Tariff (First Quote)",
             "Bar chart: blue = A1 Flat net monthly cost; green = Midday Saver net monthly cost, "
             "for the first solar+battery quote. Black step line = no-solar A1 Flat baseline. "
             "Bars below zero mean DEBS export credits exceed your import and supply charges for that month — "
             "Synergy owes you money. Summer months are cheapest because solar output is highest; "
             "winter months are more expensive but still cheaper than the baseline in most cases. "
             "This chart shows your expected monthly bill throughout the year after installation."),
            ("14_monthly_cost_verification.png",
             "Fig 5 — VERIFICATION: All Scenarios vs Modelled Baseline  [CHECK AGAINST YOUR REAL BILLS]",
             "IMPORTANT: Before trusting any savings figures, compare the solid line with circular markers "
             "(the modelled grid-only baseline) against your actual Synergy monthly bills. "
             "If the model consistently understates your bills, a separately metered circuit (hot water, "
             "pool pump on Economy tariff) is likely excluded from the smart meter data. "
             "Scale all savings proportionally: multiply by (average real bill ÷ average modelled baseline). "
             "If the match is good (within ~5%), the projections are reliable for your usage pattern. "
             "Dashed coloured lines show each solar+battery scenario for reference."),
            ("06_payback_analysis.png",
             "Fig 6 — 25-Year Cumulative Cash-Flow (All Quotes)",
             "Each line's starting point (negative y-axis) = net system cost after rebates. "
             "The line rises each year as annual savings accumulate. "
             "When a line crosses y = 0 (the horizontal black line), that is the payback year — "
             "cumulative savings have now paid off the system. "
             "After payback, every year of savings is clear profit. "
             "Solid lines = A1 Flat; dashed = Midday Saver; steeper slope = higher annual saving. "
             "Lines starting deeper in negative territory = larger, more expensive systems."),
            ("08_solar_sweep.png",
             "Fig 7 — Solar Size Optimisation (14 kWh Battery, Midday Saver)",
             f"Tests every array size from 3 kW to 20 kW with a 14 kWh battery on Midday Saver. "
             f"Left axis (orange): self-sufficiency — rises steeply at first, then flattens as the battery "
             f"fills easily and extra solar mostly exports. "
             f"Left axis (green): battery utilisation — peaks near the optimum array size. "
             f"Right axis (red dotted): annual export kWh — grows steadily with more panels. "
             f"Vertical dashed line = recommended {opt_kw:.1f} kW (the 'knee' of the self-sufficiency curve). "
             f"Beyond this size, extra panels add export income but the marginal return per dollar of panel "
             f"cost declines — payback period lengthens. Use this chart when negotiating array size with installers."),
            ("13_monthly_savings.png",
             "Fig 8 — Monthly Savings vs No-Solar Baseline (All Quotes)",
             "Monthly saving vs the no-solar baseline on the same tariff, for every quote and tariff. "
             "Positive bars = money saved that month vs operating without solar. "
             "Solid bars = A1 Flat; hatched bars = Midday Saver. "
             "Summer (Dec–Feb) shows the highest savings; winter (Jun–Aug) shows the lowest. "
             "Some winter Midday Saver bars may be lower than A1 Flat because the Midday Saver "
             "baseline is already cheaper in winter (most daytime import falls in the cheap SOP window), "
             "narrowing the room for solar to improve things further."),
            ("16_zero_net_sizing.png",
             "Fig 10 — Zero-Net Sizing: Minimum Solar Array per Battery Size",
             "Minimum solar array (kW) required for annual DEBS export credits to equal or exceed "
             "import + supply charges — a $0 net annual Synergy bill. Tariff: Midday Saver, year-0 DEBS rates. "
             "Green bars = zero-net is achievable with ≤35 kW of solar for this battery size. "
             "Red bars = not achievable within the 35 kW modelling limit. "
             "Larger batteries require less solar because they improve self-consumption and "
             "can shift late-afternoon generation to the higher-value peak DEBS window (10¢/kWh). "
             "Note: DEBS rates decline 5%/yr — size for significantly better than zero-net "
             "in year 1 to remain near zero-net in years 5–10."),
            ("09_summer_day.png",
             "Fig 11 — Representative Summer Day (December–February)",
             "Half-hourly energy flows for a typical summer day with the first solar+battery quote. "
             "Top panel: solar (yellow) peaks strongly in the middle of the day, filling the battery "
             "and exporting surplus. Battery discharges through the evening peak window (3–9 pm, red band), "
             "often eliminating grid imports entirely during that expensive period. "
             "Bottom panel: battery SoC rising sharply from mid-morning, reaching maximum by early afternoon, "
             "then steadily discharging through the evening as it powers the household. "
             "Summer days represent the best-case performance of the system."),
            ("10_autumn_day.png",
             "Fig 12 — Representative Autumn Day (March–May)",
             "Autumn shows moderate solar output — shorter days and lower sun angles than summer "
             "reduce peak generation but the battery still charges to 70–90% of capacity on most days. "
             "Evening discharge covers a meaningful share of peak demand; some grid import may occur "
             "at the tail end of the peak window when the battery runs low. "
             "Autumn and spring are often the most balanced seasons for solar+battery performance — "
             "enough sun to fill the battery, enough evening load to fully discharge it."),
            ("11_winter_day.png",
             "Fig 13 — Representative Winter Day (June–August)",
             "Perth winter is the most challenging period. The sun rises around 7 am, sets around 5:30 pm, "
             "and reaches a peak altitude of only ~35° (vs ~75° in summer), producing far less intense "
             "radiation over a shorter day. The battery typically charges to only 40–60% of capacity. "
             "Grid imports increase, particularly in the late evening. However, on Midday Saver, "
             "much of the daytime import still falls in the cheap SOP window, limiting the cost impact. "
             "The battery provides meaningful protection against the expensive peak period even when "
             "only partially charged."),
            ("12_spring_day.png",
             "Fig 14 — Representative Spring Day (September–November)",
             "Solar output recovers rapidly through spring as days lengthen and the sun climbs higher. "
             "By October the system is approaching summer performance levels. "
             "The battery returns to near-full daily cycles — charging close to capacity during the day "
             "and discharging through the evening peak. Grid imports during the peak window fall "
             "significantly compared to winter. Spring and autumn tend to be the best months for "
             "battery cycling efficiency: solar fully fills the battery AND evening demand fully empties it."),
        ]

        for fname, title, caption in chart_pages:
            fpath = out_dir / fname
            if not fpath.exists():
                continue
            img = plt.imread(str(fpath))
            fig, axes = plt.subplots(2, 1, figsize=(11, 8.5),
                                     gridspec_kw={"height_ratios": [12, 1]})
            fig.patch.set_facecolor("white")
            axes[0].imshow(img); axes[0].axis("off")
            axes[0].set_title(title, fontsize=13, fontweight="bold", pad=8)
            axes[1].axis("off")
            axes[1].text(0.5, 0.5, _t(caption), ha="center", va="center",
                         fontsize=9, color="#444444",
                         transform=axes[1].transAxes, wrap=True)
            plt.tight_layout(h_pad=0.5)
            pdf.savefig(fig, bbox_inches="tight"); plt.close(fig)

        # ── Assumptions page ─────────────────────────────────────────────────
        fig = plt.figure(figsize=(11, 8.5))
        fig.patch.set_facecolor("white")
        ax = fig.add_axes([0.06, 0.04, 0.88, 0.90]); ax.axis("off")
        ax.set_title("Key Assumptions & Methodology",
                     fontsize=16, fontweight="bold", loc="left")
        assumptions = [
            ("Tariff rates — Synergy WA, effective 1 July 2025",
             f"A1 Flat: {A1_RATE*100:.4f}¢/kWh usage rate + {A1_SUPPLY*100:.4f}¢/day supply charge (fixed). "
             f"Midday Saver: {MS_SOP*100:.4f}¢/kWh super off-peak 9am–3pm, "
             f"{MS_PEAK*100:.4f}¢/kWh peak 3–9pm, {MS_OP*100:.4f}¢/kWh off-peak 9pm–9am. "
             f"Midday Saver supply: {MS_SUPPLY*100:.4f}¢/day."),
            ("DEBS export — Distributed Energy Buyback Scheme",
             f"{DEBS_PEAK*100:.0f}¢/kWh during the peak window (3–9 pm). "
             f"{DEBS_OP*100:.0f}¢/kWh at all other times (including the cheap SOP window). "
             f"DEBS rates decline {DEBS_DECL*100:.0f}% per year — export credits become less valuable over time. "
             "Maximising self-consumption is therefore increasingly important as the analysis horizon extends."),
            ("Tariff escalation",
             f"{TARIFF_ESC*100:.1f}% per annum applied to all import charges and supply charges. "
             f"Conservative base case — Synergy has historically raised tariffs 3–6%/yr. "
             f"Savings grow each year as the electricity you avoid buying from the grid becomes more expensive. "
             f"Sensitivity: at 5%/yr escalation most payback periods shorten by approximately 1–2 years."),
            ("Battery parameters",
             f"Round-trip efficiency: {BAT_RTE*100:.0f}% (√ applied symmetrically to charge and discharge). "
             f"Depth of discharge: {BAT_DOD*100:.0f}% of nameplate capacity. "
             f"Annual capacity degradation: {BAT_DEG*100:.1f}%/yr. "
             "Dispatch priority: solar surplus charges battery first; battery discharges to cover load shortfall "
             "before grid import. Inverter output capped at rated kW per quote."),
            ("Solar modelling",
             f"Synthetic bell-curve daily profile × panel array size (kW DC) × system efficiency {SYS_EFF*100:.0f}% "
             f"(inverter + wiring losses) × seasonal Peak Sun Hours (4 hrs/day winter – 6 hrs/day summer) "
             f"× seasonal shading factor × diurnal sunrise/sunset taper. "
             f"Panel degradation {SOL_DEG*100:.2f}%/yr applied annually. "
             "Output is synthetic — does not reflect day-to-day cloud variability. Actual yield will vary."),
            ("Shading model",
             f"Seasonal factors: summer (Dec–Feb) {SHADE_SUMMER*100:.0f}%, "
             f"autumn (Mar–May) {SHADE_AUTUMN*100:.0f}%, "
             f"winter (Jun–Aug) {SHADE_WINTER*100:.0f}%, "
             f"spring (Sep–Nov) {SHADE_SPRING*100:.0f}%. "
             "Plus a diurnal taper reducing output to zero at sunrise/sunset. "
             "Combined effect ≈ 40% of theoretical unshaded output. "
             "Adjust seasonal factors upward for a largely unshaded north-facing roof (80–90% summer, 60–70% winter)."),
            ("Federal battery rebate — Cheaper Home Batteries Program (from 1 May 2026)",
             "Tiered subsidy on battery storage capacity only (not panels or inverter): "
             "$252/kWh for first 14 kWh, $151/kWh for 14–28 kWh, $38/kWh for 28–50 kWh. "
             "Capped at 40% of total installed system cost. "
             "Applied as point-of-sale discount by installer — no separate government claim required."),
            ("State battery rebate — Synergy WA Residential Battery Scheme",
             f"${STATE_REBATE_FLAT:,.0f} flat payment, requires enrolment in Synergy's Virtual Power Plant (VPP). "
             "VPP allows Synergy to occasionally dispatch a small amount of your battery capacity to stabilise "
             "the grid — typically infrequent and limited. All battery quotes assumed VPP-eligible."),
            ("Limitations and what is not modelled",
             "No NPV or discount rate — all figures are nominal AUD. No maintenance costs (inverter replacement, "
             "cleaning). No allowance for financing costs or cost of capital. "
             "Single year of consumption data applied across all 25 years. "
             "Western Power single-phase export limit (5 kW) not enforced — confirm with installer. "
             "STC solar rebate assumed already reflected in installer's gross quoted price."),
        ]
        y = 0.91
        for heading, body in assumptions:
            ax.text(0, y, _t(f"▪ {heading}:"), fontsize=9.5, fontweight="bold",
                    transform=ax.transAxes, va="top", color="#1a3a5c")
            y -= 0.038
            words = _t(body).split()
            line = ""; lines = []
            for w in words:
                if len(line) + len(w) + 1 <= 115:
                    line = (line + " " + w).strip()
                else:
                    lines.append(line); line = w
            if line: lines.append(line)
            for ln in lines:
                ax.text(0.025, y, ln, fontsize=8.5, color="#333333",
                        transform=ax.transAxes, va="top")
                y -= 0.031
            y -= 0.014
        pdf.savefig(fig, bbox_inches="tight"); plt.close(fig)

    print(f"  PDF report   → {path.resolve()}")
    return path


# ─── MAIN ─────────────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--csv",  default=None)
    parser.add_argument("--out",  default="/Users/jonathanking/Python/solar_analysis_output/results")
    parser.add_argument("--demo", action="store_true")
    args = parser.parse_args()

    print("\n"+"="*78)
    print("  SOLAR & BATTERY ECONOMIC MODELLER — Synergy WA")
    print("  Shading: seasonal+diurnal | Data: your actual 2024 meter file")
    print("="*78)

    try:
        # ── Validate static configuration first
        print("\n  Validating configuration …")
        validate_constants()
        validated_quotes = validate_quotes(QUOTES)
        _info(f"Configuration OK: {len(validated_quotes)} quotes, "
              f"{ANALYSIS_YEARS}-year horizon, "
              f"tariff escalation {TARIFF_ESC*100:.1f}% p.a.")

        # ── Load data
        print("\n  Loading consumption data …")
        if args.demo:
            raw_df = demo_data()
            raw_df = validate_consumption_data(raw_df)
            data_source = "synthetic demo data (--demo flag)"
        elif args.csv:
            raw_df = load_data(args.csv)
            data_source = str(args.csv)
        else:
            dp = Path(DEFAULT_DATA_PATH)
            if dp.exists():
                print(f"  Source: {dp}")
                raw_df = load_data(dp)
                data_source = str(dp)
            else:
                print(f"\n  ✗ DATA FILE NOT FOUND: {dp}")
                print(f"    Update DEFAULT_DATA_PATH at the top of the script to point to your")
                print(f"    Synergy half-hourly meter export file (.xlsx or .csv), or run with:")
                print(f"      python3 solar_battery_analyser.py --csv /path/to/your/file.xlsx")
                print(f"    To run a quick test with synthetic data instead, add the --demo flag.")
                sys.exit(1)

    except ValidationError as exc:
        print(f"\n  ✗ VALIDATION FAILED:\n    {exc}\n")
        sys.exit(1)

    # Pre-compute helper columns
    print("\n  Pre-computing solar profiles …")
    raw_df["slot"] = (raw_df["datetime"].dt.hour*2 + raw_df["datetime"].dt.minute//30).astype(int)
    raw_df["doy"]  = raw_df["datetime"].dt.dayofyear.astype(int)
    raw_df["date"] = raw_df["datetime"].dt.date

    out_dir = Path(args.out)
    tariffs = ["A1 Flat","Midday Saver"]
    all_res = []; pbs = []

    print(f"\n  Running {len(validated_quotes)} quotes × 2 tariffs …\n")
    for (label,solar_kw,bat_kwh,inv_kw,cost),tariff in itertools_product(validated_quotes, tariffs):
        sys.stdout.write(f"  • {label:<40} {tariff} … "); sys.stdout.flush()
        df = add_solar(raw_df, solar_kw)

        r  = simulate(df, solar_kw, bat_kwh, inv_kw, tariff)
        r["label"] = label; r["bat_kwh"] = bat_kwh
        all_res.append(r)

        pb = payback(df, solar_kw, bat_kwh, inv_kw, cost, tariff, label)
        pbs.append(pb)
        print("✓")

    print_baseline_check(raw_df, data_source)
    print_summary(all_res, pbs, raw_df)

    SWEEP_BAT_KWH = 14.0
    SWEEP_TARIFF  = "Midday Saver"
    print(f"  Sweeping solar sizes for {SWEEP_BAT_KWH} kWh battery ({SWEEP_TARIFF}) …")
    sweep_df = sweep_solar_for_battery(raw_df, bat_kwh=SWEEP_BAT_KWH, tariff=SWEEP_TARIFF)
    opt_kw   = optimal_solar_size(sweep_df)
    print_solar_sweep(sweep_df, SWEEP_BAT_KWH, opt_kw, SWEEP_TARIFF)

    # Battery-only arbitrage scenarios not configured — skipping
    arbit_res = []; arbit_pbs = []

    # ── Zero-net sizing optimisation ──────────────────────────────────────────
    ZN_TARIFF = "Midday Saver"
    zero_net  = find_zero_net_combinations(raw_df, tariff=ZN_TARIFF)
    print_zero_net(zero_net, ZN_TARIFF)

    print("  Generating charts …")
    make_dashboard(all_res, pbs, raw_df, out_dir,
                   sweep_df, SWEEP_BAT_KWH, SWEEP_TARIFF, opt_kw,
                   arbit_res, arbit_pbs, zero_net, ZN_TARIFF)

    print("  Writing outputs …")
    write_excel(all_res, pbs, raw_df, sweep_df, opt_kw, out_dir)
    write_timeseries_excel(raw_df, out_dir)
    write_report(all_res, pbs, raw_df, sweep_df, opt_kw, out_dir,
                 arbit_res, arbit_pbs, zero_net)
    write_pdf_report(all_res, pbs, raw_df, opt_kw, out_dir,
                     arbit_res, arbit_pbs, zero_net)
    print(f"\n  Done! → {out_dir.resolve()}\n")

if __name__ == "__main__":
    main()
